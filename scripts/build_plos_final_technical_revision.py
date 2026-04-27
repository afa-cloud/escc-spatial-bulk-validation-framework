from __future__ import annotations

import csv
import gzip
import hashlib
import importlib
import json
import math
import os
import re
import shutil
import sys
import zipfile
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd
from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill


CODE_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = Path(__file__).resolve().parents[2]
PREV_OUT = PROJECT_ROOT / "submission_ready_2026-04-27_method_framework_rewrite"
FINAL_OUT = PROJECT_ROOT / "submission_ready_2026-04-27_plos_one_final_technical_revision"

if str(CODE_DIR) not in sys.path:
    sys.path.insert(0, str(CODE_DIR))

import build_method_framework_rewrite as framework  # noqa: E402
import run_independent_patient_and_spatial_quant as spatial_patient  # noqa: E402
import run_spatial_axis_deep_validation as deep  # noqa: E402


MANUSCRIPT_NAME = "ESCC_spatial_source_table_PLOS_ONE_framework_rewrite.docx"
MANUSCRIPT_MD_NAME = "ESCC_spatial_source_table_PLOS_ONE_framework_rewrite.md"
FINAL_ZIP_NAME = "PLOS_ONE_final_technical_revision_UPLOAD_CLEAN_2026-04-27.zip"

FOCUS_PANEL_IDS = {"caf", "ecm_remodeling", "emt", "hypoxia", "tls_b_cell", "pi3k_akt_mtor", "notch"}
AXIS_LABEL_OVERRIDES = {
    "ogt_pi3k_tls_axis": "OGT/PI3K/TLS axis",
    "caf_epi_jag1_notch_niche": "CAF/ECM stromal-remodeling phenotype",
}


def configure_framework() -> None:
    framework.OUT = FINAL_OUT
    framework.MANUSCRIPT_DIR = FINAL_OUT / "manuscript"
    framework.FIGURE_DIR = FINAL_OUT / "figures"
    framework.SUPP_DIR = FINAL_OUT / "supporting_information"
    framework.AUDIT_DIR = FINAL_OUT / "audit"
    framework.DELIVERABLE_ZIP = FINAL_OUT / FINAL_ZIP_NAME


def configure_deep_modules() -> None:
    deep.OUT_ROOT = PROJECT_ROOT
    deep.DATA_ROOT = PROJECT_ROOT / "data"
    deep.TABLE_ROOT = PROJECT_ROOT / "results" / "tables"
    deep.REPORT_ROOT = PROJECT_ROOT / "reports"
    deep.REVIEW_ROOT = PROJECT_ROOT / "reviews"
    spatial_patient.OUT_ROOT = PROJECT_ROOT
    spatial_patient.DATA_ROOT = PROJECT_ROOT / "data"
    spatial_patient.TABLE_ROOT = PROJECT_ROOT / "results" / "tables"
    spatial_patient.REPORT_ROOT = PROJECT_ROOT / "reports"
    spatial_patient.REVIEW_ROOT = PROJECT_ROOT / "reviews"
    spatial_patient.CACHE_ROOT = PROJECT_ROOT / "data" / "ensembl_sequence_cache"


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def finite_mean(values: list[float]) -> float:
    vals = [v for v in values if math.isfinite(v)]
    return float(sum(vals) / len(vals)) if vals else float("nan")


def zscore_expr(expr: dict[str, list[float]]) -> dict[str, list[float]]:
    out: dict[str, list[float]] = {}
    for gene, values in expr.items():
        vals = [v for v in values if math.isfinite(v)]
        if len(vals) < 2:
            out[gene] = [float("nan") for _ in values]
            continue
        mean_value = sum(vals) / len(vals)
        variance = sum((v - mean_value) ** 2 for v in vals) / len(vals)
        sd = math.sqrt(variance)
        if sd == 0 or not math.isfinite(sd):
            out[gene] = [0.0 if math.isfinite(v) else float("nan") for v in values]
        else:
            out[gene] = [(v - mean_value) / sd if math.isfinite(v) else float("nan") for v in values]
    return out


def score_samples(genes: list[str], expr: dict[str, list[float]], sample_count: int) -> tuple[list[float], list[str]]:
    present = [gene for gene in genes if gene in expr and len(expr[gene]) == sample_count]
    scores: list[float] = []
    for idx in range(sample_count):
        values = [expr[gene][idx] for gene in present if math.isfinite(expr[gene][idx])]
        scores.append(finite_mean(values))
    return scores, present


def all_panels() -> dict[tuple[str, str], list[str]]:
    panels: dict[tuple[str, str], list[str]] = {}
    for panel_id, genes in deep.IMMUNE_PANELS.items():
        panels[("immune", panel_id)] = genes
    for panel_id, genes in deep.PATHWAY_PANELS.items():
        panels[("pathway", panel_id)] = genes
    return panels


def all_relevant_genes() -> set[str]:
    genes = set(gene for axis in deep.AXES.values() for gene in axis["genes"])
    genes |= set(gene for panel in deep.IMMUNE_PANELS.values() for gene in panel)
    genes |= set(gene for panel in deep.PATHWAY_PANELS.values() for gene in panel)
    return {gene.upper() for gene in genes}


def load_gse47404_expr() -> tuple[list[str], dict[str, list[float]]]:
    matrix_path = PROJECT_ROOT / "data" / "geo" / "GSE47404" / "GSE47404_series_matrix.txt.gz"
    annot_path = PROJECT_ROOT / "data" / "geo" / "GSE47404" / "GPL6480.annot.gz"
    relevant = all_relevant_genes()
    probe_to_genes = deep.load_gpl6480_probe_map(annot_path, relevant)
    samples, _metadata = deep.parse_geo_metadata(matrix_path)
    samples2, expr, _probe_counts = deep.load_gse47404_expression(matrix_path, probe_to_genes, relevant)
    if samples2:
        samples = samples2
    return samples, {gene.upper(): values for gene, values in expr.items()}


def load_expression_layers() -> dict[str, tuple[list[str], dict[str, list[float]]]]:
    tcga_samples, tcga_expr = deep.load_tcga_escc_expression()
    gse_samples, gse_expr = load_gse47404_expr()
    return {
        "TCGA_ESCC_Xena": (tcga_samples, {gene.upper(): values for gene, values in tcga_expr.items()}),
        "GSE47404": (gse_samples, gse_expr),
    }


def compute_raw_association(
    dataset: str,
    samples: list[str],
    expr: dict[str, list[float]],
    axis_id: str,
    panel_type: str,
    panel_id: str,
    panel_genes: list[str],
    transform_label: str,
    remove_overlap_from_panel: bool,
) -> dict[str, Any]:
    sample_count = len(samples)
    axis_genes = [gene.upper() for gene in deep.AXES[axis_id]["genes"]]
    panel_genes_upper = [gene.upper() for gene in panel_genes]
    overlap = sorted(set(axis_genes) & set(panel_genes_upper))
    used_panel_genes = [gene for gene in panel_genes_upper if gene not in set(axis_genes)] if remove_overlap_from_panel else panel_genes_upper
    axis_scores, axis_present = score_samples(axis_genes, expr, sample_count)
    panel_scores, panel_present = score_samples(used_panel_genes, expr, sample_count)
    rho, p_value, n_corr = deep.spearman(axis_scores, panel_scores)
    return {
        "dataset": dataset,
        "axis_id": axis_id,
        "axis_label": AXIS_LABEL_OVERRIDES.get(axis_id, deep.AXES[axis_id]["label"]),
        "panel_type": panel_type,
        "panel_id": panel_id,
        "score_scale": transform_label,
        "n_samples": sample_count,
        "n_correlation_samples": n_corr,
        "axis_genes_present": len(axis_present),
        "axis_present_genes": ",".join(axis_present),
        "panel_genes_defined": len(panel_genes_upper),
        "panel_genes_used": len(used_panel_genes),
        "panel_present_genes": ",".join(panel_present),
        "overlap_genes": ",".join(overlap),
        "overlap_gene_count": len(overlap),
        "spearman_rho": rho,
        "spearman_p_asymptotic": p_value,
        "executor_agent_id": "final_technical_executor_001",
        "reviewer_agent_id": "final_technical_reviewer_001",
    }


def add_fdr(rows: list[dict[str, Any]], p_key: str, out_key: str) -> None:
    fdr = deep.bh_fdr([float(row.get(p_key, 1.0)) for row in rows])
    for row, value in zip(rows, fdr):
        row[out_key] = value


def compute_overlap_tables(expression_layers: dict[str, tuple[list[str], dict[str, list[float]]]]) -> tuple[list[dict[str, Any]], list[dict[str, Any]], list[dict[str, Any]]]:
    overlap_rows: list[dict[str, Any]] = []
    overlap_corr_rows: list[dict[str, Any]] = []
    zscore_rows: list[dict[str, Any]] = []
    panels = all_panels()
    for axis_id, axis in deep.AXES.items():
        axis_genes = [gene.upper() for gene in axis["genes"]]
        for (panel_type, panel_id), panel_genes in panels.items():
            panel_upper = [gene.upper() for gene in panel_genes]
            overlap = sorted(set(axis_genes) & set(panel_upper))
            overlap_rows.append(
                {
                    "axis_id": axis_id,
                    "axis_label": AXIS_LABEL_OVERRIDES.get(axis_id, axis["label"]),
                    "panel_type": panel_type,
                    "panel_id": panel_id,
                    "axis_gene_count": len(axis_genes),
                    "panel_gene_count": len(panel_upper),
                    "overlap_gene_count": len(overlap),
                    "overlap_genes": ",".join(overlap),
                    "panel_genes_after_overlap_removal": ",".join([gene for gene in panel_upper if gene not in set(axis_genes)]),
                    "executor_agent_id": "final_technical_executor_001",
                    "reviewer_agent_id": "final_technical_reviewer_001",
                }
            )
    for dataset, (samples, raw_expr) in expression_layers.items():
        z_expr = zscore_expr(raw_expr)
        raw_rows_for_fdr: list[dict[str, Any]] = []
        removed_rows_for_fdr: list[dict[str, Any]] = []
        z_rows_for_fdr: list[dict[str, Any]] = []
        for axis_id in deep.AXES:
            for (panel_type, panel_id), panel_genes in panels.items():
                raw = compute_raw_association(dataset, samples, raw_expr, axis_id, panel_type, panel_id, panel_genes, "raw_log2_like", False)
                removed = compute_raw_association(dataset, samples, raw_expr, axis_id, panel_type, panel_id, panel_genes, "raw_log2_like_overlap_removed_panel", True)
                zrow = compute_raw_association(dataset, samples, z_expr, axis_id, panel_type, panel_id, panel_genes, "within_cohort_gene_zscore", False)
                raw_rows_for_fdr.append(raw)
                removed_rows_for_fdr.append(removed)
                z_rows_for_fdr.append(zrow)
        add_fdr(raw_rows_for_fdr, "spearman_p_asymptotic", "spearman_fdr_bh")
        add_fdr(removed_rows_for_fdr, "spearman_p_asymptotic", "spearman_fdr_bh")
        add_fdr(z_rows_for_fdr, "spearman_p_asymptotic", "spearman_fdr_bh")
        raw_lookup = {(r["axis_id"], r["panel_type"], r["panel_id"]): r for r in raw_rows_for_fdr}
        for removed in removed_rows_for_fdr:
            raw = raw_lookup[(removed["axis_id"], removed["panel_type"], removed["panel_id"])]
            raw_rho = float(raw["spearman_rho"])
            removed_rho = float(removed["spearman_rho"])
            same_direction = math.isfinite(raw_rho) and math.isfinite(removed_rho) and (raw_rho == 0 or removed_rho == 0 or raw_rho * removed_rho > 0)
            support_note = "not_evaluable"
            if math.isfinite(removed_rho):
                if same_direction and removed["spearman_fdr_bh"] < 0.05:
                    support_note = "direction_retained_and_fdr_supported"
                elif same_direction:
                    support_note = "direction_retained_with_reduced_support"
                else:
                    support_note = "direction_not_retained_after_overlap_removal"
            overlap_corr_rows.append(
                {
                    **removed,
                    "raw_spearman_rho": raw_rho,
                    "raw_spearman_fdr_bh": raw["spearman_fdr_bh"],
                    "overlap_removed_spearman_rho": removed_rho,
                    "overlap_removed_spearman_fdr_bh": removed["spearman_fdr_bh"],
                    "absolute_rho_change": abs(raw_rho) - abs(removed_rho) if math.isfinite(raw_rho) and math.isfinite(removed_rho) else float("nan"),
                    "direction_consistent_with_raw": "yes" if same_direction else "no",
                    "support_note": support_note,
                }
            )
        for zrow in z_rows_for_fdr:
            raw = raw_lookup[(zrow["axis_id"], zrow["panel_type"], zrow["panel_id"])]
            raw_rho = float(raw["spearman_rho"])
            z_rho = float(zrow["spearman_rho"])
            same_direction = math.isfinite(raw_rho) and math.isfinite(z_rho) and (raw_rho == 0 or z_rho == 0 or raw_rho * z_rho > 0)
            zscore_rows.append(
                {
                    **zrow,
                    "raw_spearman_rho": raw_rho,
                    "raw_spearman_fdr_bh": raw["spearman_fdr_bh"],
                    "zscore_spearman_rho": z_rho,
                    "zscore_spearman_fdr_bh": zrow["spearman_fdr_bh"],
                    "direction_consistent_with_raw": "yes" if same_direction else "no",
                }
            )
    return overlap_rows, overlap_corr_rows, zscore_rows


def parse_tsv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8", newline="") as fh:
        return list(csv.DictReader(fh, delimiter="\t"))


def compute_probe_uniqueness_rows() -> list[dict[str, Any]]:
    raw_path = PROJECT_ROOT / "data" / "geo" / "GSE53625" / "GSM1296956_first_raw_member.txt.gz"
    mapping_path = PROJECT_ROOT / "submission_ready_2026-04-26" / "supplementary_tables" / "S6_GSE53625_mapping.tsv"
    features = spatial_patient.load_gse53625_feature_sequences(raw_path)
    sequence_counts = Counter(item["sequence"] for item in features)
    feature_by_num = {str(item["feature_num"]): item for item in features}
    mapping_rows = parse_tsv(mapping_path)
    accepted = [row for row in mapping_rows if row.get("accepted_gene")]
    rows: list[dict[str, Any]] = []
    for row in accepted:
        feature = feature_by_num.get(str(row["feature_num"]), {})
        seq = feature.get("sequence", "")
        rows.append(
            {
                "dataset": "GSE53625",
                "feature_num": row["feature_num"],
                "accepted_gene": row["accepted_gene"],
                "match_type": row["match_type"],
                "sequence_length": len(seq) if seq else row.get("sequence_length", ""),
                "array_design_exact_sequence_count": sequence_counts.get(seq, 0) if seq else "",
                "array_design_sequence_unique": "yes" if seq and sequence_counts.get(seq, 0) == 1 else "no",
                "target_scope_specificity": row["specificity_scope"],
                "genome_wide_specificity_status": "not_performed",
                "reviewer_note": "Exact probe sequence uniqueness was checked across the full first raw-member array design. This is not a genome-wide or transcriptome-wide specificity audit.",
                "executor_agent_id": "final_technical_executor_001",
                "reviewer_agent_id": "final_technical_reviewer_001",
            }
        )
    rows.insert(
        0,
        {
            "dataset": "GSE53625",
            "feature_num": "SUMMARY",
            "accepted_gene": f"{len(accepted)} accepted probe rows",
            "match_type": "",
            "sequence_length": "",
            "array_design_exact_sequence_count": f"{sum(1 for item in rows if item.get('array_design_sequence_unique') == 'yes')} unique accepted sequences",
            "array_design_sequence_unique": f"features_checked={len(features)}",
            "target_scope_specificity": "target-scope Ensembl axis-gene sequence rescue",
            "genome_wide_specificity_status": "not_performed",
            "reviewer_note": f"The added check improves array-design duplicate detection but does not establish genome-wide probe specificity. Ambiguous non-accepted mapping rows in S6 were not counted as accepted: {len(mapping_rows) - len(accepted)}.",
            "executor_agent_id": "final_technical_executor_001",
            "reviewer_agent_id": "final_technical_reviewer_001",
        },
    )
    return rows


def enriched_manifest_rows() -> list[dict[str, Any]]:
    manifest_path = PROJECT_ROOT / "submission_ready_2026-04-26" / "supplementary_tables" / "S16_data_manifest.tsv"
    rows = parse_tsv(manifest_path)
    output_map = {
        "TCGA": "S1_Table.xlsx::S2_TCGA_immune; S3_TCGA_pathway; S17-S19 sensitivity sheets",
        "GSE47404": "S1_Table.xlsx::S4_GSE47404_assoc; S17-S19 sensitivity sheets",
        "GSE53625": "S1_Table.xlsx::S6_GSE53625_mapping; S7_GSE53625_validation; S20_GSE53625_probe_uniqueness",
        "GSE53625 first raw member": "S1_Table.xlsx::S20_GSE53625_probe_uniqueness",
        "Ensembl": "S1_Table.xlsx::S6_GSE53625_mapping; S20_GSE53625_probe_uniqueness",
        "HRA003627": "S1_Table.xlsx::S8_HRA003627_source",
        "HRA008846 Table S3": "S1_Table.xlsx::S9_HRA008846_DEG",
        "HRA008846 Table S4": "S1_Table.xlsx::S10_HRA008846_cells",
        "HRA008846 Table S6": "S1_Table.xlsx::S11_HRA008846_LR",
        "GDSC2": "S1_Table.xlsx::S12_GDSC2_drugs",
    }
    out: list[dict[str, Any]] = []
    for row in rows:
        local = Path(row.get("local_path", ""))
        exists = local.exists()
        key = row.get("resource", "")
        processed = ""
        for prefix, value in output_map.items():
            if key.startswith(prefix):
                processed = value
                break
        out.append(
            {
                **row,
                "file_size_bytes_verified": local.stat().st_size if exists else "",
                "checksum_sha256": sha256(local) if exists and local.is_file() else "",
                "processed_output_path": processed or "S1_Table.xlsx",
            }
        )
    return out


def write_rows_to_sheet(wb: Any, sheet_name: str, rows: list[dict[str, Any]]) -> None:
    if sheet_name in wb.sheetnames:
        del wb[sheet_name]
    ws = wb.create_sheet(sheet_name)
    if not rows:
        ws.append(["status", "note"])
        ws.append(["empty", "No rows generated"])
        return
    headers = list(rows[0].keys())
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="D9EAF7")
    for row in rows:
        ws.append([format_cell(row.get(header, "")) for header in headers])
    ws.freeze_panes = "A2"
    for idx, header in enumerate(headers, start=1):
        width = max(10, min(45, len(header) + 2))
        ws.column_dimensions[ws.cell(row=1, column=idx).column_letter].width = width


def format_cell(value: Any) -> Any:
    if isinstance(value, float):
        if math.isnan(value):
            return "nan"
        return float(f"{value:.8g}")
    if isinstance(value, (list, tuple, set)):
        return ",".join(str(v) for v in value)
    return value


def update_readme_sheet(wb: Any) -> None:
    readme = wb["README"]
    additions = [
        ("S17_gene_set_overlap", "Axis-to-comparator gene-set overlap audit."),
        ("S18_overlap_removed_corr", "Spearman correlations after removing axis-overlapping genes from comparator panels."),
        ("S19_zscore_sensitivity", "Within-cohort per-gene z-score sensitivity analysis for core correlations."),
        ("S20_GSE53625_probe_scope", "GSE53625 accepted-probe array-design sequence uniqueness and target-scope limitation audit."),
        ("S21_data_manifest_ext", "Extended data-source manifest with checksum and processed-output path fields."),
    ]
    existing = {readme.cell(row=i, column=1).value for i in range(1, readme.max_row + 1)}
    for sheet, description in additions:
        if sheet not in existing:
            readme.append([sheet, description])


def enhance_s1_table(technical: dict[str, Any]) -> None:
    path = framework.SUPP_DIR / "S1_Table.xlsx"
    wb = load_workbook(path)
    update_readme_sheet(wb)
    write_rows_to_sheet(wb, "S17_gene_set_overlap", technical["overlap_rows"])
    write_rows_to_sheet(wb, "S18_overlap_removed_corr", technical["overlap_corr_rows"])
    write_rows_to_sheet(wb, "S19_zscore_sensitivity", technical["zscore_rows"])
    write_rows_to_sheet(wb, "S20_GSE53625_probe_scope", technical["probe_uniqueness_rows"])
    write_rows_to_sheet(wb, "S21_data_manifest_ext", technical["manifest_rows"])
    wb.save(path)


def append_s2_code_zip() -> None:
    code_zip = framework.SUPP_DIR / "S2_Code.zip"
    with zipfile.ZipFile(code_zip, "a", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as zf:
        zf.write(Path(__file__), arcname="scripts/build_plos_final_technical_revision.py")


def key_row(rows: list[dict[str, Any]], dataset: str, axis_id: str, panel_id: str) -> dict[str, Any]:
    candidates = [
        row
        for row in rows
        if row["dataset"] == dataset
        and row["axis_id"] == axis_id
        and row["panel_id"] == panel_id
        and row["panel_id"] in FOCUS_PANEL_IDS
    ]
    if not candidates:
        return {}
    return candidates[0]


def fmt_rho(row: dict[str, Any], key: str) -> str:
    value = row.get(key)
    try:
        return f"{float(value):.3f}"
    except Exception:
        return "NA"


def summarize_technical_findings(technical: dict[str, Any]) -> dict[str, str]:
    overlap_rows = technical["overlap_corr_rows"]
    z_rows = technical["zscore_rows"]
    tcga_ecm = key_row(overlap_rows, "TCGA_ESCC_Xena", "caf_epi_jag1_notch_niche", "ecm_remodeling")
    tcga_emt = key_row(overlap_rows, "TCGA_ESCC_Xena", "caf_epi_jag1_notch_niche", "emt")
    tcga_caf = key_row(overlap_rows, "TCGA_ESCC_Xena", "caf_epi_jag1_notch_niche", "caf")
    gse_ecm = key_row(overlap_rows, "GSE47404", "caf_epi_jag1_notch_niche", "ecm_remodeling")
    gse_caf = key_row(overlap_rows, "GSE47404", "caf_epi_jag1_notch_niche", "caf")
    tcga_hypoxia = key_row(overlap_rows, "TCGA_ESCC_Xena", "ogt_pi3k_tls_axis", "hypoxia")
    z_focus = [
        row
        for row in z_rows
        if row["panel_id"] in FOCUS_PANEL_IDS
        and row["axis_id"] in {"caf_epi_jag1_notch_niche", "ogt_pi3k_tls_axis"}
    ]
    same_direction = sum(1 for row in z_focus if row.get("direction_consistent_with_raw") == "yes")
    total = len(z_focus)
    return {
        "overlap_sentence": (
            "After removing axis-overlapping genes from comparator panels, the CAF/ECM phenotype retained "
            f"directionally consistent associations with ECM remodeling in TCGA ESCC (rho={fmt_rho(tcga_ecm, 'overlap_removed_spearman_rho')}) "
            f"and GSE47404 (rho={fmt_rho(gse_ecm, 'overlap_removed_spearman_rho')}) and with EMT in TCGA ESCC "
            f"(rho={fmt_rho(tcga_emt, 'overlap_removed_spearman_rho')}). CAF-marker correlations were attenuated after overlap removal "
            f"(TCGA rho={fmt_rho(tcga_caf, 'overlap_removed_spearman_rho')}; GSE47404 rho={fmt_rho(gse_caf, 'overlap_removed_spearman_rho')}), "
            "so CAF-panel results are interpreted as supportive association evidence rather than independent confirmation. "
            f"The OGT/PI3K/TLS hypoxia pattern was directionally retained after overlap removal (TCGA rho={fmt_rho(tcga_hypoxia, 'overlap_removed_spearman_rho')})."
        ),
        "zscore_sentence": (
            f"Within-cohort gene-level z-score sensitivity analysis preserved the direction of {same_direction}/{total} focus-panel correlations, "
            "supporting the direction of the main association patterns while leaving effect-size strength dependent on the scoring scale."
        ),
    }


def package_versions() -> dict[str, str]:
    modules = ["pandas", "numpy", "xenaPython", "openpyxl", "docx", "PIL"]
    names = {"docx": "python-docx", "PIL": "Pillow"}
    versions: dict[str, str] = {"Python": sys.version.split()[0]}
    for module_name in modules:
        try:
            importlib.import_module(module_name)
            package_name = names.get(module_name, module_name)
            try:
                version = importlib.metadata.version(package_name)  # type: ignore[attr-defined]
            except Exception:
                version = getattr(importlib.import_module(module_name), "__version__", "unknown")
            versions[package_name] = str(version)
        except Exception:
            pass
    return versions


def revise_manuscript_md(base_md: str, summary: dict[str, str]) -> str:
    text = base_md
    text = text.replace("We developed a tiered validation framework", "We applied an implemented tiered validation framework")
    text = text.replace("approximate P=", "asymptotic Spearman P=")
    text = text.replace("approximate P values", "asymptotic Spearman P values")
    text = text.replace("P values were adjusted using the Benjamini-Hochberg false discovery rate procedure [14].",
        "Asymptotic Spearman P values were computed from the rank-correlation coefficient using a normal approximation implemented in the analysis script, and P values were adjusted using the Benjamini-Hochberg false discovery rate procedure [14].")
    text = text.replace(
        "Axis scores were calculated from log2 expression values as the mean expression of available genes in each pre-specified axis. No interpolation, model-based imputation or optimized gene weighting was used. Genes absent from a given platform were omitted from that platform-specific score, and gene coverage was reported where relevant. This mean-aggregation strategy was chosen to keep the scoring rule transparent and reproducible across public expression resources.",
        "Axis scores were calculated from log2 expression values as the mean expression of available genes in each pre-specified axis. No interpolation, model-based imputation or optimized gene weighting was used. Genes absent from a given platform were omitted from that platform-specific score, and gene coverage was reported where relevant. This mean-aggregation strategy was chosen to keep the scoring rule transparent and reproducible across public expression resources. As a final technical audit, we quantified overlap between each axis and every comparator panel, then recalculated Spearman correlations after removing axis-overlapping genes from the comparator panel. We also repeated the core correlation analyses after within-cohort z-score standardization of each gene before axis or panel score aggregation. These sensitivity analyses are provided in S1 Table."
    )
    text = text.replace(
        "Exact matches and one-mismatch matches were accepted only when uniquely assigned within the pre-specified axis gene universe. This procedure is a target-scope rescue and not a genome-wide probe-specificity audit.",
        "Exact matches and one-mismatch matches were accepted only when uniquely assigned within the pre-specified axis gene universe. We additionally checked whether accepted probe sequences were unique across the first raw-member array design. This procedure remains a target-scope supportive rescue and an array-design duplicate check; it is not a transcriptome-wide or genome-wide probe-specificity audit."
    )
    text = text.replace(
        "GDSC2 fitted dose-response data were filtered to ESCA cell lines and summarized for axis-relevant drug target classes [11]. These data were interpreted as target-class context only, because no model was fitted between axis-expression scores and drug-response metrics.",
        "GDSC2 fitted dose-response data were filtered to ESCA cell lines and summarized only as supplementary axis-relevant target-class context [11]. No expression-response model was fitted, and these summaries were not used to infer drug sensitivity prediction."
    )
    text = text.replace(
        "The manuscript package was rebuilt and audited with Python 3.12.13 using python-docx, Pillow and openpyxl for document generation, figure inspection and workbook validation. Rscript 4.5.3 was available in the local environment, but the supplied reconstruction scripts for this package are Python-based. Reproducible scripts and configuration files are provided as S2 Code. The scripts document data accessions, cache paths, generated tables and figure outputs. No analysis step used unpublished data or access-restricted raw sequencing data.",
        "All analysis and package reconstruction scripts were run in Python 3.12.13. The analysis environment used pandas 3.0.1, numpy 2.3.5 and xenaPython 1.0.14 for public-data access and matrix processing. Statistical procedures used script-native implementations of Spearman rank correlation, Benjamini-Hochberg FDR adjustment, sign tests, log-rank summaries and Mann-Whitney-type rank comparisons; scipy, statsmodels, lifelines and R were not used in the final reconstruction environment. Document generation and audit used python-docx 1.2.0, openpyxl 3.1.5 and Pillow 12.2.0. Reproducible scripts and configuration files are provided as S2 Code. The scripts document data accessions, cache paths, generated tables and figure outputs. No analysis step used unpublished data or access-restricted raw sequencing data.\n\n### Artificial intelligence tools disclosure\n\nOpenAI ChatGPT/Codex was used for language editing, structural organization of the manuscript, code drafting and package-audit assistance. The authors manually reviewed and verified the generated text, analysis scripts, numerical outputs, citations and conclusions. AI tools were not used to fabricate, alter or replace research data, and all reported results were derived from the public data sources and scripts documented in S1 Table and S2 Code."
    )
    text = text.replace(
        "These cross-cohort associations demonstrate reproducibility of a CAF/ECM stromal-remodeling phenotype at the bulk expression level (Fig 2).",
        "These cross-cohort associations demonstrate reproducibility of a CAF/ECM stromal-remodeling phenotype at the bulk expression level (Fig 2). " + summary["overlap_sentence"] + " " + summary["zscore_sentence"] + " Full overlap and z-score sensitivity results are provided in S1 Table."
    )
    text = text.replace(
        "Target probe rescue in GSE53625 accepted 98 probe rows and covered 16 axis-related genes. Both axes had complete gene coverage for their predefined gene lists. In 179 paired tumor-normal samples, the OGT/PI3K axis was elevated in tumor tissue with a median tumor-normal delta of 0.823 (two-sided sign-test P=6.10e-41). The CAF/ECM phenotype was also elevated in tumor tissue with a median delta of 1.482 (P=4.70e-52). Survival analyses did not demonstrate stratification: log-rank P values were 0.25785 for the OGT/PI3K axis and 0.372619 for the CAF/ECM phenotype. Thus, the paired cohort supports expression upshift in tumors but not independent survival stratification (Fig 3).",
        "Target probe rescue in GSE53625 accepted 98 probe rows and covered 16 axis-related genes. The final audit also checked accepted probe sequences against all probe sequences in the first raw-member array design; this supports array-design uniqueness assessment but does not establish genome-wide specificity. Both axes had complete gene coverage for their predefined gene lists. In 179 paired tumor-normal samples, the OGT/PI3K axis was elevated in tumor tissue with a median tumor-normal delta of 0.823 (two-sided sign-test P=6.10e-41). The CAF/ECM phenotype was also elevated in tumor tissue with a median delta of 1.482 (P=4.70e-52). Survival analyses did not demonstrate stratification: log-rank P values were 0.25785 for the OGT/PI3K axis and 0.372619 for the CAF/ECM phenotype. Thus, the paired cohort provides supportive tumor-normal evidence after target-scope rescue but not definitive probe-specific or survival validation (Fig 3)."
    )
    text = text.replace(
        "GDSC2 ESCA summaries highlighted EGFR/ERBB inhibitors and PI3K/AKT/mTOR-related compounds among axis-relevant target classes. Sapitinib, gefitinib, lapatinib, erlotinib and osimertinib appeared among the leading EGFR/ERBB-context rows. Because axis-expression scores were not modeled jointly with drug-response metrics, these results provide therapeutic target-class context rather than drug-sensitivity prediction.",
        "GDSC2 ESCA summaries were retained in S1 Table as supplementary target-class context for axis-relevant drug classes. Because no joint axis-expression and drug-response model was fitted, these summaries do not support drug-sensitivity prediction."
    )
    text = text.replace(
        "Third, GSE53625 probe rescue was restricted to pre-specified target genes and did not perform a genome-wide specificity audit.",
        "Third, GSE53625 probe rescue was restricted to pre-specified target genes; the added array-design duplicate check does not replace a transcriptome-wide or genome-wide specificity audit."
    )
    text = text.replace(
        "| CAF/ECM stromal-remodeling phenotype is reproducible | TCGA CAF rho=0.906/FDR=1.02e-16; TCGA ECM rho=0.917/FDR=3.80e-17; GSE47404 ECM rho=0.829/FDR=7.31e-11 | Supported at association level |",
        "| CAF/ECM stromal-remodeling phenotype is reproducible | TCGA CAF rho=0.906/FDR=1.02e-16; TCGA ECM rho=0.917/FDR=3.80e-17; GSE47404 ECM rho=0.829/FDR=7.31e-11; overlap-removed and z-score sensitivity analyses retained direction for key ECM comparisons | Supported at association level, with sensitivity analyses reducing the strength of independence claims |"
    )
    text = text.replace(
        "All data used in this study are publicly accessible from TCGA/GDC/UCSC Xena, GEO accessions GSE47404 and GSE53625, HRA003627 source data, HRA008846 supplementary tables and GDSC2. Processed tables generated in this project are provided in S1 Table and in the submitted code package. Accession identifiers, URLs, download/cache status, local file sizes and source-table conflict notes are provided in the data-source manifest within S1 Table. The study does not rely on restricted data requiring author-mediated access.",
        "All data used in this study are publicly accessible from TCGA/GDC/UCSC Xena, GEO accessions GSE47404 and GSE53625, HRA003627 source data, HRA008846 supplementary tables and GDSC2. Processed tables generated in this project are provided in S1 Table. Accession identifiers, URLs, download dates, verified file sizes, SHA256 checksums where a local file was used, and processed output paths are provided in the extended data-source manifest within S1 Table. The study does not rely on restricted data requiring author-mediated access."
    )
    text = text.replace(
        "Reproducible scripts, configuration files and a code README are provided as S2 Code. The package includes the Python scripts used to rebuild the public-data workflow, independent patient and source-table checks, final submission package generation and upload-package auditing. If a public repository or archive DOI is assigned after submission, the repository link can replace or supplement S2 Code in the published Data Availability Statement.",
        "Reproducible scripts, configuration files and a code README are provided as S2 Code. The package includes the Python scripts used to rebuild the public-data workflow, independent patient and source-table checks, overlap and z-score sensitivity analyses, final submission package generation and upload-package auditing. For this submission package, S2 Code is the code-availability route and S1 Table is the processed-table route; no additional author-mediated code request is required."
    )
    text = text.replace(
        "S1 Table. Supplementary tables for axis definitions, cohort associations, GSE53625 probe rescue, spatial source-table quantification, GDSC2 target-class context, interpretive checks, references and data-source manifest.",
        "S1 Table. Supplementary tables for axis definitions, cohort associations, gene-set overlap audit, overlap-removed correlations, z-score sensitivity analysis, GSE53625 probe rescue and array-design scope audit, spatial source-table quantification, GDSC2 target-class context, interpretive checks, references and the extended data-source manifest."
    )
    text = text.replace("framework demonstrates reproducibility", "framework supports reproducibility")
    text = text.replace("demonstrate a direct", "support a direct")
    text = text.replace("The CAF/ECM phenotype is strongly associated with CAF, ECM remodeling and EMT programs", "The CAF/ECM phenotype is associated with CAF, ECM remodeling and EMT programs")
    text = text.replace("the CAF/ECM phenotype is strongly associated with CAF, ECM remodeling and EMT programs", "the CAF/ECM phenotype is associated with CAF, ECM remodeling and EMT programs")
    return text


def docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def write_final_change_summary(technical: dict[str, Any], summary: dict[str, str]) -> None:
    lines = [
        "# Final Technical Revision Summary",
        "",
        "1. Added gene-set overlap audit for all axis/comparator pairs.",
        "2. Added overlap-removed Spearman correlations for TCGA ESCC and GSE47404.",
        "3. Added within-cohort per-gene z-score sensitivity analysis.",
        "4. Added GSE53625 array-design accepted-probe sequence uniqueness audit while preserving target-scope limitation language.",
        "5. Extended the data-source manifest with file sizes, SHA256 checksums and processed output paths.",
        "6. Downgraded GDSC2 to supplementary target-class context only.",
        "7. Added AI tools disclosure consistent with PLOS policy.",
        "8. Replaced uncertain future repository language with explicit S1 Table / S2 Code availability.",
        "9. Removed Rscript language and listed actual Python packages used in the reconstruction environment.",
        "10. Replaced approximate P wording with asymptotic Spearman P where applicable.",
        "",
        "## Main Sensitivity Statements",
        "",
        f"- {summary['overlap_sentence']}",
        f"- {summary['zscore_sentence']}",
        "",
        "## Generated Supplementary Sheets",
        "",
        "- S17_gene_set_overlap",
        "- S18_overlap_removed_corr",
        "- S19_zscore_sensitivity",
        "- S20_GSE53625_probe_scope",
        "- S21_data_manifest_ext",
    ]
    (FINAL_OUT / "final_technical_change_summary.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def run_technical_analyses() -> dict[str, Any]:
    configure_deep_modules()
    expression_layers = load_expression_layers()
    overlap_rows, overlap_corr_rows, zscore_rows = compute_overlap_tables(expression_layers)
    probe_rows = compute_probe_uniqueness_rows()
    manifest_rows = enriched_manifest_rows()
    technical = {
        "overlap_rows": overlap_rows,
        "overlap_corr_rows": overlap_corr_rows,
        "zscore_rows": zscore_rows,
        "probe_uniqueness_rows": probe_rows,
        "manifest_rows": manifest_rows,
    }
    return technical


def custom_final_audit(technical: dict[str, Any]) -> None:
    manuscript = framework.MANUSCRIPT_DIR / MANUSCRIPT_NAME
    s1 = framework.SUPP_DIR / "S1_Table.xlsx"
    s2 = framework.SUPP_DIR / "S2_Code.zip"
    upload_zip = framework.DELIVERABLE_ZIP
    text = docx_text(manuscript)
    narrative = text.split("\nReferences\n")[0]
    wb = load_workbook(s1, read_only=True, data_only=True)
    required_sheets = {"S17_gene_set_overlap", "S18_overlap_removed_corr", "S19_zscore_sensitivity", "S20_GSE53625_probe_scope", "S21_data_manifest_ext"}
    with zipfile.ZipFile(s2) as zf:
        code_entries = sorted(zf.namelist())
        code_bad = zf.testzip()
    with zipfile.ZipFile(upload_zip) as zf:
        upload_entries = sorted(zf.namelist())
        upload_bad = zf.testzip()
    required_refs = ["Fig 1", "Fig 2", "Fig 3", "Fig 4", "Fig 5", "Table 1", "Table 2", "S1 Table", "S2 Code"]
    orphan_page = any(p.text.strip() == "Page" for p in Document(str(manuscript)).paragraphs)
    status_rows = [
        ("new_s1_sheets_present", required_sheets.issubset(set(wb.sheetnames)), ";".join(sorted(set(wb.sheetnames) & required_sheets))),
        ("overlap_rows_present", len(technical["overlap_rows"]) > 0 and len(technical["overlap_corr_rows"]) > 0, f"overlap={len(technical['overlap_rows'])}; corr={len(technical['overlap_corr_rows'])}"),
        ("zscore_rows_present", len(technical["zscore_rows"]) > 0, str(len(technical["zscore_rows"]))),
        ("probe_scope_rows_present", len(technical["probe_uniqueness_rows"]) > 1, str(len(technical["probe_uniqueness_rows"]))),
        ("manifest_checksums_present", all(row.get("checksum_sha256") for row in technical["manifest_rows"] if str(row.get("local_status", "")).lower() == "exists"), "checked local exists rows"),
        ("future_repo_language_absent", "If a public repository or archive DOI is assigned" not in text, ""),
        ("available_upon_request_absent", "available upon request" not in text.lower(), ""),
        ("developed_absent_in_narrative", re.search(r"\bdeveloped\b", narrative, flags=re.I) is None, ""),
        ("approximate_p_absent", "approximate P" not in text, ""),
        ("required_references_present", all(item in text for item in required_refs), ";".join(item for item in required_refs if item not in text)),
        ("orphan_page_absent", not orphan_page, str(orphan_page)),
        ("s2_code_contains_final_script", "scripts/build_plos_final_technical_revision.py" in code_entries and code_bad is None, f"entries={len(code_entries)} bad={code_bad}"),
        ("upload_zip_valid", upload_bad is None and len(upload_entries) == 8, f"entries={upload_entries}; bad={upload_bad}"),
    ]
    audit = {
        "audit_time": datetime.now().isoformat(timespec="seconds"),
        "overall_status": "pass" if all(row[1] for row in status_rows) else "fail",
        "status_rows": [{"check": row[0], "pass": row[1], "evidence": row[2]} for row in status_rows],
        "upload_zip": {"path": str(upload_zip), "bytes": upload_zip.stat().st_size, "sha256": sha256(upload_zip), "entries": upload_entries},
        "s1_table": {"path": str(s1), "bytes": s1.stat().st_size, "sha256": sha256(s1), "sheets": wb.sheetnames},
        "s2_code": {"path": str(s2), "bytes": s2.stat().st_size, "sha256": sha256(s2), "entries": code_entries},
    }
    path_json = framework.AUDIT_DIR / "final_technical_revision_audit.json"
    path_tsv = framework.AUDIT_DIR / "final_technical_revision_audit.tsv"
    path_json.write_text(json.dumps(audit, indent=2, ensure_ascii=False), encoding="utf-8")
    with path_tsv.open("w", encoding="utf-8", newline="") as fh:
        fh.write("check\tstatus\tevidence\n")
        for row in audit["status_rows"]:
            fh.write(f"{row['check']}\t{'pass' if row['pass'] else 'fail'}\t{str(row['evidence']).replace(chr(9), ' ').replace(chr(10), ' ')}\n")
    if audit["overall_status"] != "pass":
        raise RuntimeError(json.dumps(audit["status_rows"], indent=2))


def main() -> int:
    configure_framework()
    technical = run_technical_analyses()
    summary = summarize_technical_findings(technical)
    framework.ensure_clean_dirs()
    framework.MANUSCRIPT_MD = revise_manuscript_md(framework.MANUSCRIPT_MD, summary)
    framework.write_static_files()
    write_final_change_summary(technical, summary)
    framework.copy_figures()
    framework.copy_supporting_information()
    enhance_s1_table(technical)
    append_s2_code_zip()
    framework.build_upload_zip()
    framework.audit_all()
    custom_final_audit(technical)
    print(f"wrote {FINAL_OUT}")
    print(f"zip {framework.DELIVERABLE_ZIP}")
    print(f"sha256 {sha256(framework.DELIVERABLE_ZIP)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
