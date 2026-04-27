from __future__ import annotations

import hashlib
import json
import re
import shutil
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import load_workbook
from PIL import Image, ImageChops


BASE = Path(__file__).resolve().parents[1]
SOURCE_PLOS_DIR = BASE / "plos_one_submission"
PLOS_DIR = BASE / "plos_one_upload_clean_2026-04-26"
DELIVERABLES = BASE.parent / "deliverables"
AUDIT_DIR = BASE / "plos_one_upload_audit"
ZIP_OUT = DELIVERABLES / "ESCC_spatial_source_table_PLOS_ONE_UPLOAD_CLEAN_2026-04-26.zip"

EXPECTED_UPLOAD_FILES = [
    "manuscript/ESCC_spatial_source_table_PLOS_ONE_main_text.docx",
    "figures/Fig1.tif",
    "figures/Fig2.tif",
    "figures/Fig3.tif",
    "figures/Fig4.tif",
    "figures/Fig5.tif",
    "supporting_information/S1_Table.xlsx",
]


def prepare_upload_dir() -> None:
    resolved_base = BASE.resolve()
    resolved_upload_parent = PLOS_DIR.parent.resolve()
    if resolved_upload_parent != resolved_base or PLOS_DIR.name != "plos_one_upload_clean_2026-04-26":
        raise RuntimeError(f"Refusing to rebuild unexpected upload directory: {PLOS_DIR}")
    if PLOS_DIR.exists():
        shutil.rmtree(PLOS_DIR)
    for rel_path in EXPECTED_UPLOAD_FILES:
        (PLOS_DIR / rel_path).parent.mkdir(parents=True, exist_ok=True)
    for rel_path in EXPECTED_UPLOAD_FILES:
        src = SOURCE_PLOS_DIR / rel_path
        dst = PLOS_DIR / rel_path
        if rel_path.endswith(".docx"):
            doc = Document(str(src))
            for section in doc.sections:
                sectPr = section._sectPr
                for old in list(sectPr.findall(qn("w:lnNumType"))):
                    sectPr.remove(old)
                ln = OxmlElement("w:lnNumType")
                ln.set(qn("w:countBy"), "1")
                ln.set(qn("w:restart"), "continuous")
                sectPr.append(ln)
            doc.save(str(dst))
        else:
            shutil.copy2(src, dst)


def rel(path: Path) -> str:
    return path.relative_to(PLOS_DIR).as_posix()


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def add_check(rows: list[dict[str, Any]], file: str, check: str, status: str, evidence: str) -> None:
    rows.append({"file": file, "check": check, "status": status, "evidence": evidence})


def audit_folder(rows: list[dict[str, Any]]) -> dict[str, Any]:
    actual = sorted(rel(p) for p in PLOS_DIR.rglob("*") if p.is_file())
    expected = sorted(EXPECTED_UPLOAD_FILES)
    missing = [p for p in expected if p not in actual]
    extra = [p for p in actual if p not in expected]
    add_check(rows, ".", "expected file whitelist", "pass" if not missing and not extra else "fail",
              f"expected={len(expected)} actual={len(actual)} missing={missing} extra={extra}")
    return {"actual": actual, "expected": expected, "missing": missing, "extra": extra}


def docx_text(document: Document) -> str:
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append("\t".join(cells))
    return "\n".join(parts)


def audit_docx(path: Path, rows: list[dict[str, Any]]) -> dict[str, Any]:
    info: dict[str, Any] = {"path": rel(path), "bytes": path.stat().st_size, "sha256": sha256(path)}
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
        names = zf.namelist()
        media = [n for n in names if n.startswith("word/media/")]
        comments = [n for n in names if n.startswith("word/comments")]
        settings_xml = zf.read("word/settings.xml").decode("utf-8", errors="ignore") if "word/settings.xml" in names else ""
        footer_xml = "\n".join(
            zf.read(n).decode("utf-8", errors="ignore") for n in names if n.startswith("word/footer") and n.endswith(".xml")
        )
        document_xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    document = Document(str(path))
    text = docx_text(document)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    tables = len(document.tables)
    fig_caption_presence = {f"Fig {i}": bool(re.search(rf"\bFig\.?\s*{i}\b", text)) for i in range(1, 6)}
    references = [p for p in paragraphs if re.match(r"^(\[\d+\]|\d+\.)\s+", p)]
    placeholder_regex = re.compile(
        r"TODO|TBD|XXX|PLACEHOLDER|INSERT\s+HERE|AUTHOR\s+NAME|AFFILIATION\s+HERE|EMAIL\s+HERE|"
        r"to be added by the study team|should be completed by the study team|to be completed|to be added|"
        r"\[to be completed\]|\[to be added\]|\[author|\[affiliation|\[funding",
        re.IGNORECASE,
    )
    placeholders = sorted(set(m.group(0) for m in placeholder_regex.finditer(text)))
    title_page_text = "\n".join(paragraphs[:6])
    email_present = bool(re.search(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}", title_page_text))
    affiliation_present = bool(re.search(
        r"\b(affiliation|department|school|college|university|hospital|institute|center|centre|laboratory)\b",
        title_page_text,
        re.IGNORECASE,
    ))
    track_change_hits = len(re.findall(r"<w:(ins|del)\b", document_xml))
    line_numbering = "lnNumType" in settings_xml or "lnNumType" in document_xml
    page_field = "PAGE" in footer_xml or "PAGE" in document_xml
    info.update({
        "zip_test_bad_entry": bad,
        "paragraph_count": len(paragraphs),
        "table_count": tables,
        "embedded_media_count": len(media),
        "comment_part_count": len(comments),
        "track_change_hit_count": track_change_hits,
        "line_numbering_xml": line_numbering,
        "page_field_present": page_field,
        "fig_caption_presence": fig_caption_presence,
        "reference_count": len(references),
        "placeholder_hits": placeholders,
        "first_paragraphs": paragraphs[:5],
    })
    add_check(rows, rel(path), "valid docx zip", "pass" if bad is None else "fail", f"testzip_bad_entry={bad}")
    add_check(rows, rel(path), "no embedded figure media in main manuscript", "pass" if not media else "fail",
              f"embedded_media_count={len(media)} media={media[:10]}")
    add_check(rows, rel(path), "no comments or tracked changes", "pass" if not comments and track_change_hits == 0 else "fail",
              f"comment_parts={comments} track_change_hits={track_change_hits}")
    add_check(rows, rel(path), "line and page numbering present", "pass" if line_numbering and page_field else "warn",
              f"line_numbering_xml={line_numbering} page_field_present={page_field}")
    add_check(rows, rel(path), "caption coverage Fig 1-5", "pass" if all(fig_caption_presence.values()) else "fail",
              json.dumps(fig_caption_presence, ensure_ascii=False))
    add_check(rows, rel(path), "tables and references present", "pass" if tables >= 2 and len(references) >= 10 else "warn",
              f"tables={tables} numbered_references={len(references)}")
    add_check(rows, rel(path), "placeholder scan", "pass" if not placeholders else "warn",
              f"placeholder_hits={placeholders}")
    add_check(rows, rel(path), "corresponding author email present", "pass" if email_present else "warn",
              f"email_present={email_present}")
    add_check(rows, rel(path), "author affiliation metadata present", "pass" if affiliation_present else "warn",
              f"affiliation_present={affiliation_present}")
    return info


def audit_tif(path: Path, rows: list[dict[str, Any]]) -> dict[str, Any]:
    with Image.open(path) as im:
        width, height = im.size
        mode = im.mode
        raw_dpi = im.info.get("dpi", ("", ""))
        if isinstance(raw_dpi, tuple):
            dpi = tuple(float(x) if x else 0.0 for x in raw_dpi[:2])
        else:
            dpi = (0.0, 0.0)
        rgb = im.convert("RGB")
        diff = ImageChops.difference(rgb, Image.new("RGB", rgb.size, (255, 255, 255)))
        bbox = diff.getbbox()
    min_dpi = min(float(x) for x in dpi[:2]) if isinstance(dpi, tuple) and len(dpi) >= 2 and all(x for x in dpi[:2]) else 0
    info = {
        "path": rel(path),
        "bytes": path.stat().st_size,
        "sha256": sha256(path),
        "width_px": width,
        "height_px": height,
        "mode": mode,
        "dpi": dpi,
        "nonwhite_bbox": bbox,
    }
    add_check(rows, rel(path), "valid tif dimensions", "pass" if width >= 1800 and height >= 1000 else "fail",
              f"width={width} height={height} mode={mode}")
    add_check(rows, rel(path), "300 dpi metadata", "pass" if 295 <= min_dpi <= 305 else "warn",
              f"dpi={dpi}")
    add_check(rows, rel(path), "nonblank image content", "pass" if bbox else "fail", f"nonwhite_bbox={bbox}")
    return info


def audit_xlsx(path: Path, rows: list[dict[str, Any]]) -> dict[str, Any]:
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
    wb = load_workbook(path, read_only=True, data_only=False)
    sheet_infos = []
    total_nonempty = 0
    for ws in wb.worksheets:
        nonempty = 0
        first_row = []
        for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
            values = [v for v in row if v is not None and str(v) != ""]
            nonempty += len(values)
            if row_idx == 1:
                first_row = [str(v) if v is not None else "" for v in row[:12]]
        total_nonempty += nonempty
        sheet_infos.append({
            "name": ws.title,
            "state": ws.sheet_state,
            "max_row": ws.max_row,
            "max_column": ws.max_column,
            "nonempty_cells": nonempty,
            "first_row": first_row,
        })
    info = {
        "path": rel(path),
        "bytes": path.stat().st_size,
        "sha256": sha256(path),
        "zip_test_bad_entry": bad,
        "sheet_count": len(sheet_infos),
        "sheets": sheet_infos,
        "external_links": len(getattr(wb, "_external_links", [])),
    }
    add_check(rows, rel(path), "valid xlsx zip", "pass" if bad is None else "fail", f"testzip_bad_entry={bad}")
    add_check(rows, rel(path), "nonempty workbook", "pass" if sheet_infos and total_nonempty > 0 else "fail",
              f"sheets={len(sheet_infos)} total_nonempty_cells={total_nonempty}")
    hidden = [s["name"] for s in sheet_infos if s["state"] != "visible"]
    add_check(rows, rel(path), "no hidden sheets", "pass" if not hidden else "warn", f"hidden_sheets={hidden}")
    add_check(rows, rel(path), "no external workbook links", "pass" if info["external_links"] == 0 else "warn",
              f"external_links={info['external_links']}")
    return info


def build_zip(rows: list[dict[str, Any]]) -> dict[str, Any]:
    DELIVERABLES.mkdir(parents=True, exist_ok=True)
    if ZIP_OUT.exists():
        ZIP_OUT.unlink()
    with zipfile.ZipFile(ZIP_OUT, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as zf:
        for rel_path in EXPECTED_UPLOAD_FILES:
            zf.write(PLOS_DIR / rel_path, arcname=rel_path)
    with zipfile.ZipFile(ZIP_OUT) as zf:
        bad = zf.testzip()
        entries = sorted(zf.namelist())
        missing = [p for p in EXPECTED_UPLOAD_FILES if p not in entries]
        extra = [p for p in entries if p not in EXPECTED_UPLOAD_FILES]
        hash_mismatches = []
        entry_summaries = []
        for rel_path in EXPECTED_UPLOAD_FILES:
            source = PLOS_DIR / rel_path
            data = zf.read(rel_path)
            entry_hash = sha256_bytes(data)
            source_hash = sha256(source)
            if entry_hash != source_hash:
                hash_mismatches.append(rel_path)
            entry_summaries.append({
                "path": rel_path,
                "zip_bytes": len(data),
                "source_bytes": source.stat().st_size,
                "sha256": entry_hash,
            })
    status = "pass" if bad is None and not missing and not extra and not hash_mismatches else "fail"
    add_check(rows, ZIP_OUT.name, "zip exact whitelist and byte hash match", status,
              f"testzip_bad_entry={bad} missing={missing} extra={extra} hash_mismatches={hash_mismatches}")
    return {
        "path": str(ZIP_OUT),
        "bytes": ZIP_OUT.stat().st_size,
        "sha256": sha256(ZIP_OUT),
        "entries": entry_summaries,
        "missing": missing,
        "extra": extra,
        "hash_mismatches": hash_mismatches,
        "zip_test_bad_entry": bad,
    }


def write_reports(payload: dict[str, Any], rows: list[dict[str, Any]]) -> None:
    AUDIT_DIR.mkdir(parents=True, exist_ok=True)
    (AUDIT_DIR / "plos_one_upload_content_audit.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    with (AUDIT_DIR / "plos_one_upload_content_audit.tsv").open("w", encoding="utf-8", newline="") as fh:
        fh.write("file\tcheck\tstatus\tevidence\n")
        for row in rows:
            evidence = str(row["evidence"]).replace("\t", " ").replace("\r", " ").replace("\n", " ")
            fh.write(f"{row['file']}\t{row['check']}\t{row['status']}\t{evidence}\n")
    manifest_rows = []
    for rel_path in EXPECTED_UPLOAD_FILES:
        p = PLOS_DIR / rel_path
        manifest_rows.append({
            "path": rel_path,
            "bytes": p.stat().st_size,
            "sha256": sha256(p),
        })
    with (AUDIT_DIR / "plos_one_upload_manifest.tsv").open("w", encoding="utf-8", newline="") as fh:
        fh.write("path\tbytes\tsha256\n")
        for row in manifest_rows:
            fh.write(f"{row['path']}\t{row['bytes']}\t{row['sha256']}\n")
    md = [
        "# PLOS ONE Clean Upload Audit",
        "",
        f"Audit time: {payload['audit_time']}",
        f"Clean upload zip: `{payload['zip']['path']}`",
        f"Zip SHA256: `{payload['zip']['sha256']}`",
        "",
        "## File Verdicts",
        "",
        "| File | Check | Status | Evidence |",
        "|---|---|---|---|",
    ]
    for row in rows:
        evidence = str(row["evidence"]).replace("|", "\\|").replace("\n", " ")
        md.append(f"| {row['file']} | {row['check']} | {row['status']} | {evidence} |")
    (AUDIT_DIR / "PLOS_ONE_UPLOAD_AUDIT.md").write_text("\n".join(md) + "\n", encoding="utf-8")


def main() -> int:
    rows: list[dict[str, Any]] = []
    prepare_upload_dir()
    folder_info = audit_folder(rows)
    if folder_info["missing"] or folder_info["extra"]:
        payload = {
            "audit_time": datetime.now().isoformat(timespec="seconds"),
            "folder": folder_info,
            "error": "Upload folder does not match expected whitelist.",
        }
        write_reports(payload, rows)
        print(json.dumps(payload, ensure_ascii=False, indent=2))
        return 2
    file_audits: list[dict[str, Any]] = []
    for rel_path in EXPECTED_UPLOAD_FILES:
        path = PLOS_DIR / rel_path
        if path.suffix.lower() == ".docx":
            file_audits.append(audit_docx(path, rows))
        elif path.suffix.lower() in {".tif", ".tiff"}:
            file_audits.append(audit_tif(path, rows))
        elif path.suffix.lower() == ".xlsx":
            file_audits.append(audit_xlsx(path, rows))
        else:
            file_audits.append({"path": rel_path, "bytes": path.stat().st_size, "sha256": sha256(path)})
    tif_hashes = [a["sha256"] for a in file_audits if a["path"].startswith("figures/")]
    add_check(rows, "figures", "no duplicate TIFF byte content", "pass" if len(set(tif_hashes)) == len(tif_hashes) else "fail",
              f"tif_count={len(tif_hashes)} unique_hashes={len(set(tif_hashes))}")
    zip_info = build_zip(rows)
    failed = [r for r in rows if r["status"] == "fail"]
    warned = [r for r in rows if r["status"] == "warn"]
    payload = {
        "audit_time": datetime.now().isoformat(timespec="seconds"),
        "folder": folder_info,
        "files": file_audits,
        "zip": zip_info,
        "failed_checks": failed,
        "warning_checks": warned,
        "overall_status": "pass" if not failed else "fail",
    }
    write_reports(payload, rows)
    print(json.dumps(payload, ensure_ascii=False, indent=2))
    return 0 if not failed else 1


if __name__ == "__main__":
    raise SystemExit(main())
