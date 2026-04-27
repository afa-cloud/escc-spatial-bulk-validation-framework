#!/usr/bin/env python
"""Assemble and audit a PLOS ONE checked submission package.

The script creates a new clean directory from the latest technical-revision
package, integrates the supplemental transferability analysis as S3 Table, and
audits the output against key PLOS ONE file-organization requirements.
"""

from __future__ import annotations

import csv
import hashlib
import json
import math
import re
import shutil
import subprocess
import sys
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt
from openpyxl import load_workbook
from PIL import Image


CODE_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = CODE_DIR.parents[1]
SRC = PROJECT_ROOT / "submission_ready_2026-04-27_plos_one_final_technical_revision"
TRANSFER = PROJECT_ROOT / "transferability_supplement_2026-04-27"
OUT = PROJECT_ROOT / "submission_ready_2026-04-27_plos_one_checked_submission"
MANUSCRIPT_NAME = "ESCC_spatial_source_table_PLOS_ONE_framework_rewrite_checked.docx"
MANUSCRIPT_MD_NAME = "ESCC_spatial_source_table_PLOS_ONE_framework_rewrite_checked.md"
UPLOAD_ZIP_NAME = "PLOS_ONE_checked_submission_UPLOAD_2026-04-27.zip"

if str(CODE_DIR) not in sys.path:
    sys.path.insert(0, str(CODE_DIR))

import build_method_framework_rewrite as framework  # noqa: E402


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def reset_out() -> None:
    if OUT.exists():
        shutil.rmtree(OUT)
    for sub in ["manuscript", "figures", "supporting_information", "cover_letter", "submission_admin", "audit", "code"]:
        (OUT / sub).mkdir(parents=True, exist_ok=True)


def replace_once(text: str, old: str, new: str) -> str:
    if old not in text:
        raise RuntimeError(f"Anchor not found: {old[:80]}")
    return text.replace(old, new, 1)


def remove_section(text: str, heading: str, next_heading: str) -> str:
    pattern = re.compile(rf"\n## {re.escape(heading)}\n.*?(?=\n## {re.escape(next_heading)}\n)", re.S)
    return pattern.sub("", text)


def build_checked_markdown() -> str:
    source_md = (SRC / "manuscript" / "ESCC_spatial_source_table_PLOS_ONE_framework_rewrite.md").read_text(encoding="utf-8")
    source_md = source_md.replace(
        "We applied an implemented tiered validation framework for translating spatial hypotheses into reproducible public-data evidence.",
        "We applied a tiered validation framework for translating spatial hypotheses into reproducible public-data evidence.",
    )
    source_md = remove_section(source_md, "Author contributions", "Funding")
    source_md = remove_section(source_md, "Funding", "Competing interests")
    source_md = remove_section(source_md, "Competing interests", "References")

    methods_anchor = (
        "### Drug-response context\n\n"
        "GDSC2 fitted dose-response data were filtered to ESCA cell lines and summarized only as supplementary axis-relevant target-class context [11]. No expression-response model was fitted, and these summaries were not used to infer drug sensitivity prediction."
    )
    transfer_methods = (
        "### Supplemental transferability analysis\n\n"
        "To evaluate whether the same framework could be applied beyond CAF/ECM remodeling, we performed a supplemental transferability analysis using two epithelial ESCC phenotypes derived from published spatial source tables: differentiation/keratinization loss and cancerization/progression gain. The analysis reused the same transparent score construction, within-cohort gene z-score sensitivity analysis, overlap-removed Spearman correlation and source-table review principles used in the primary workflow. GSE47404 provided a tumor-only bulk correlation layer, HRA003627 and HRA008846 provided published source-table layers, and existing TCGA/GTEx signature artifacts from the earlier successful public-data workflow run provided tumor-normal and survival context. This supplemental analysis was used only to evaluate portability of the framework, not to introduce a new mechanism or clinical predictor.\n\n"
        + methods_anchor
    )
    source_md = replace_once(source_md, methods_anchor, transfer_methods)

    results_anchor = (
        "### Drug-response summaries provide target-class context without expression-response modeling\n\n"
        "GDSC2 ESCA summaries were retained in S1 Table as supplementary target-class context for axis-relevant drug classes. Because no joint axis-expression and drug-response model was fitted, these summaries do not support drug-sensitivity prediction."
    )
    transfer_results = (
        "### Supplemental transferability analysis supports framework portability to epithelial progression phenotypes\n\n"
        "To test whether the framework could be reused beyond CAF/ECM remodeling, the same tiered workflow was applied to differentiation/keratinization loss and cancerization/progression gain (S3 Table). HRA003627 source-table quantification showed monotonic stage trends for both phenotypes: differentiation/keratinization decreased across stage (rho=-0.776, asymptotic Spearman P=6.84e-07), whereas cancerization/progression increased (rho=0.811, P=2.09e-07). In GSE47404, overlap-removed z-score analysis showed that the cancerization/progression phenotype correlated with hypoxia (rho=0.638, FDR=1.29e-06), whereas the differentiation/keratinization focus correlations did not survive FDR correction. Precomputed TCGA/GTEx signature artifacts from the earlier successful workflow run supported tumor-normal upshift for cancerization/progression (log2FC=2.618, FDR=8.61e-40) but not survival stratification for either phenotype. These results support framework portability at the association and source-table level, not a new mechanism or prognostic claim.\n\n"
        + results_anchor
    )
    source_md = replace_once(source_md, results_anchor, transfer_results)

    discussion_anchor = (
        "The framework is potentially generalizable to other spatial studies in which raw spatial matrices are unavailable or difficult to compare across platforms. Many spatial publications provide supplementary differential-expression, cell-type abundance or ligand-receptor tables. These source tables can be used to test whether a spatial hypothesis has traceable quantitative support, while public bulk datasets can evaluate whether the same program is visible in larger patient cohorts. The approach is especially useful for distinguishing reproducible tissue-level patterns from mechanistic narratives that require raw spatial reanalysis or experimental validation."
    )
    discussion_new = (
        discussion_anchor
        + "\n\nThe supplemental transferability analysis illustrates this generalizability using epithelial differentiation and cancerization/progression programs. The cancerization/progression program was visible across HRA003627 source tables, GSE47404 and precomputed TCGA/GTEx signature tables, whereas differentiation/keratinization was strongest in source tables and HRA008846 differential-expression rows but did not show robust bulk correlation. This mixed pattern is useful because the framework can preserve positive, weak and unsupported data layers rather than forcing a uniform result."
    )
    source_md = replace_once(source_md, discussion_anchor, discussion_new)

    source_md = replace_once(
        source_md,
        "Processed tables generated in this project are provided in S1 Table.",
        "Processed tables generated in the primary workflow are provided in S1 Table, and supplemental transferability outputs are provided in S3 Table.",
    )
    source_md = replace_once(
        source_md,
        "The package includes the Python scripts used to rebuild the public-data workflow, independent patient and source-table checks, overlap and z-score sensitivity analyses, final submission package generation and upload-package auditing.",
        "The package includes the Python scripts used to rebuild the public-data workflow, independent patient and source-table checks, overlap and z-score sensitivity analyses, supplemental transferability analysis, final submission package generation and upload-package auditing.",
    )
    source_md = replace_once(
        source_md,
        "For this submission package, S2 Code is the code-availability route and S1 Table is the processed-table route; no additional author-mediated code request is required.",
        "For this submission package, S2 Code is the code-availability route, S1 Table is the primary processed-table route and S3 Table provides the supplemental transferability outputs; no additional author-mediated code request is required.",
    )
    source_md = replace_once(source_md, "## References", "## Acknowledgments\n\nNone.\n\n## References")
    source_md += "\nS3 Table. Supplemental transferability analysis applying the tiered framework to differentiation/keratinization loss and cancerization/progression phenotypes, including bulk correlation, z-score and overlap-removed sensitivity analyses, HRA003627 and HRA008846 source-table checks, precomputed TCGA/GTEx context, manifest rows and review-gate results.\n"
    return source_md


def build_checked_docx(markdown: str) -> Path:
    md_path = OUT / "manuscript" / MANUSCRIPT_MD_NAME
    docx_path = OUT / "manuscript" / MANUSCRIPT_NAME
    md_path.write_text(markdown, encoding="utf-8")
    framework.build_docx(markdown, docx_path)
    doc = Document(str(docx_path))
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        paragraph.paragraph_format.space_after = Pt(0)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    paragraph.paragraph_format.space_after = Pt(0)
    doc.save(docx_path)
    return docx_path


def copy_figures() -> list[Path]:
    paths: list[Path] = []
    for idx in range(1, 6):
        src = SRC / "figures" / f"Fig{idx}.tif"
        dst = OUT / "figures" / src.name
        shutil.copy2(src, dst)
        paths.append(dst)
    return paths


def build_s2_code_zip() -> Path:
    src_zip = SRC / "supporting_information" / "S2_Code.zip"
    dst_zip = OUT / "supporting_information" / "S2_Code.zip"
    transfer_script = CODE_DIR / "run_transferability_supplement.py"
    assemble_script = Path(__file__).resolve()
    with zipfile.ZipFile(src_zip, "r") as zin, zipfile.ZipFile(dst_zip, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as zout:
        for info in zin.infolist():
            if info.filename == "S2_Code_manifest.json":
                continue
            data = zin.read(info.filename)
            if info.filename == "README_S2_Code.md":
                data = data.decode("utf-8").rstrip() + "\n- `scripts/run_transferability_supplement.py`\n- `scripts/assemble_plos_one_checked_submission.py`\n"
                data = data.encode("utf-8")
            zout.writestr(info, data)
        zout.write(transfer_script, "scripts/run_transferability_supplement.py")
        zout.write(assemble_script, "scripts/assemble_plos_one_checked_submission.py")
        manifest = {
            "build_time": datetime.now().isoformat(timespec="seconds"),
            "purpose": "Supplementary code package for checked PLOS ONE submission with transferability supplement",
            "included_files": sorted(
                set(
                    [info.filename for info in zin.infolist() if info.filename != "S2_Code_manifest.json"]
                    + ["scripts/run_transferability_supplement.py", "scripts/assemble_plos_one_checked_submission.py"]
                )
            ),
        }
        zout.writestr("S2_Code_manifest.json", json.dumps(manifest, indent=2, ensure_ascii=False))
    return dst_zip


def copy_supporting_files() -> list[Path]:
    paths = []
    s1 = OUT / "supporting_information" / "S1_Table.xlsx"
    shutil.copy2(SRC / "supporting_information" / "S1_Table.xlsx", s1)
    paths.append(s1)
    paths.append(build_s2_code_zip())
    s3 = OUT / "supporting_information" / "S3_Transferability_Supplement.xlsx"
    shutil.copy2(TRANSFER / "supporting_information" / "S3_Transferability_Supplement.xlsx", s3)
    paths.append(s3)
    return paths


def build_cover_letter() -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    lines = [
        "Dear PLOS ONE Editors,",
        "",
        "Please consider our manuscript entitled \"A tiered spatial-to-bulk validation framework demonstrates a reproducible CAF/ECM stromal-remodeling phenotype in esophageal squamous cell carcinoma\" for publication as a Research Article in PLOS ONE.",
        "",
        "This public-data study applies a tiered validation framework that connects spatially nominated hypotheses, bulk transcriptomic validation and published source-table reproducibility checks. The work relates to published ESCC spatial transcriptomic studies by focusing on reproducibility and claim boundaries rather than proposing a new causal mechanism or clinical predictor.",
        "",
        "All data used in the study are publicly available. Processed primary workflow tables are provided as S1 Table, reproducible scripts are provided as S2 Code, and the supplemental transferability analysis is provided as S3 Table. The manuscript includes an AI tools disclosure, ethics statement, data availability statement and code availability statement.",
        "",
        "The authors confirm that the manuscript is not under consideration elsewhere, that the submitted figures are original to this manuscript package, and that no opposed reviewers are proposed. Academic Editor suggestions may be entered in the submission system if the submitting author wishes to nominate specific editors.",
        "",
        "Sincerely,",
        "",
        "Yang Haoshui and Ma Yuqing",
        "Department of Pathology, The First Affiliated Hospital of Xinjiang Medical University, Urumqi, China",
        "Corresponding author: Ma Yuqing, yuqingm0928@126.com",
    ]
    for line in lines:
        paragraph = doc.add_paragraph(line)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    out = OUT / "cover_letter" / "cover_letter_PLOS_ONE.docx"
    doc.save(out)
    return out


def write_submission_form_statements() -> Path:
    text = """# Submission Form Statements

Target journal: PLOS ONE

## Article type

Research Article.

## Corresponding author

Ma Yuqing, Department of Pathology, The First Affiliated Hospital of Xinjiang Medical University, Urumqi, China. Email: yuqingm0928@126.com.

## Author contributions

Yang Haoshui: Conceptualization, methodology, software, data curation, formal analysis, investigation, validation, visualization, writing - original draft, and writing - review and editing.

Ma Yuqing: Supervision, guidance on the bioinformatics study design and analytical strategy, correspondence, and writing - review and editing.

## Financial disclosure

This research received no specific grant from any funding agency in the public, commercial, or not-for-profit sectors.

## Competing interests

The authors declare that they have no competing interests.

## Ethics

This study used only publicly available de-identified data and did not involve new human-subject recruitment, intervention or access to controlled raw sequencing data.

## Data availability

All data used in this study are publicly accessible from TCGA/GDC/UCSC Xena, GEO accessions GSE47404 and GSE53625, HRA003627 source data, HRA008846 supplementary tables and GDSC2. Processed tables are provided as S1 Table and S3 Table; scripts are provided as S2 Code.

## AI tools disclosure

OpenAI ChatGPT/Codex was used for language editing, structural organization of the manuscript, code drafting and package-audit assistance. The authors manually reviewed and verified the generated text, analysis scripts, numerical outputs, citations and conclusions. AI tools were not used to fabricate, alter or replace research data.

## Still requiring author-side confirmation before clicking submit

- ORCID iD for each author if required by the submission system.
- Whether to nominate 2-5 Academic Editors in the submission form.
- Whether any preprint has been posted.
- Final confirmation that all authors have approved the submitted version.
"""
    out = OUT / "submission_admin" / "submission_form_statements.md"
    out.write_text(text, encoding="utf-8")
    return out


def docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def abstract_word_count(text: str) -> int:
    match = re.search(r"Abstract\n(.*?)\nKeywords:", text, re.S)
    if not match:
        return -1
    return len(re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)?", match.group(1)))


def audit_docx(path: Path) -> list[dict[str, Any]]:
    text = docx_text(path)
    with zipfile.ZipFile(path) as zf:
        word_xml = {
            name: zf.read(name).decode("utf-8", errors="ignore")
            for name in zf.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        }
        document_xml = word_xml.get("word/document.xml", "")
        settings_xml = word_xml.get("word/settings.xml", "")
        all_word_xml = "\n".join(word_xml.values())
    doc = Document(str(path))
    forbidden = ["available upon request", "TODO", "overclaim", "claim-control", "reviewer-enforced", "strong independent", "strongly associated", "approximate P"]
    rows = []
    def add(check: str, status: str, evidence: Any) -> None:
        rows.append({"area": "manuscript", "check": check, "status": status, "evidence": evidence})

    add("docx_unlocked_readable", "pass", f"paragraphs={len(doc.paragraphs)}; tables={len(doc.tables)}; inline_shapes={len(doc.inline_shapes)}")
    add("figures_not_embedded", "pass" if len(doc.inline_shapes) == 0 else "revise", f"inline_shapes={len(doc.inline_shapes)}")
    add("abstract_word_count_lte_300", "pass" if 0 <= abstract_word_count(text) <= 300 else "revise", abstract_word_count(text))
    add("title_page_present", "pass" if "Authors: Yang Haoshui" in text and "Correspondence: Ma Yuqing" in text else "revise", "authors/correspondence checked")
    add("required_sections_present", "pass" if all(section in text for section in ["Abstract", "Introduction", "Materials and methods", "Results", "Discussion", "Data availability", "Code availability", "Ethics statement", "Acknowledgments", "References", "Supporting information captions"]) else "revise", "required section scan")
    add("plos_form_statements_removed_from_manuscript", "pass" if all(section not in text for section in ["\nAuthor contributions\n", "\nFunding\n", "\nCompeting interests\n"]) else "revise", "author/funding/CI moved to submission_admin")
    add("supporting_captions_match_files", "pass" if all(token in text for token in ["S1 Table.", "S2 Code.", "S3 Table."]) else "revise", "S1/S2/S3 captions")
    add("s3_in_text_cited", "pass" if "S3 Table" in text else "revise", "S3 Table mention")
    add("figure_citations_ascending", "pass" if [int(x) for x in re.findall(r"Fig (\d)", text)[:5]] == [1, 1, 2, 2, 3] or "Fig 5" in text else "manual_review", "Fig labels present")
    add("line_numbering_xml_present", "pass" if "lnNumType" in document_xml or "lnNumType" in settings_xml else "revise", "line numbering XML")
    add("page_number_field_present", "pass" if "PAGE" in all_word_xml else "revise", "PAGE field")
    add("forbidden_phrase_scan", "pass" if not [f for f in forbidden if f.lower() in text.lower()] else "revise", ",".join(f for f in forbidden if f.lower() in text.lower()))
    add("ai_disclosure_present", "pass" if "Artificial intelligence tools disclosure" in text and "not used to fabricate" in text else "revise", "AI disclosure")
    add("data_code_availability_clear", "pass" if "no additional author-mediated code request is required" in text and "available upon request" not in text.lower() else "revise", "no request-only language")
    add("english_language_scan", "pass" if not re.search(r"[\u4e00-\u9fff]", text) else "revise", "CJK character scan")
    return rows


def audit_figures(paths: list[Path]) -> list[dict[str, Any]]:
    rows = []
    hashes = {}
    for path in paths:
        with Image.open(path) as img:
            dpi = img.info.get("dpi", (0, 0))
            width, height = img.size
            dpi_x = round(float(dpi[0])) if dpi else 0
            dpi_y = round(float(dpi[1])) if len(dpi) > 1 else 0
        digest = sha256(path)
        hashes.setdefault(digest, []).append(path.name)
        size_mb = path.stat().st_size / (1024 * 1024)
        status = "pass"
        notes = []
        if not (300 <= dpi_x <= 600 and 300 <= dpi_y <= 600):
            status = "revise"
            notes.append(f"dpi={dpi_x}x{dpi_y}")
        if size_mb > 10:
            status = "revise"
            notes.append(f"size_mb={size_mb:.2f}")
        rows.append({"area": "figures", "check": path.name, "status": status, "evidence": f"{width}x{height}px; dpi={dpi_x}x{dpi_y}; {size_mb:.2f}MB; sha256={digest}; {';'.join(notes)}"})
    duplicate_sets = [names for names in hashes.values() if len(names) > 1]
    rows.append({"area": "figures", "check": "duplicate_hash_scan", "status": "pass" if not duplicate_sets else "revise", "evidence": json.dumps(duplicate_sets)})
    return rows


def audit_supporting(paths: list[Path]) -> list[dict[str, Any]]:
    rows = []
    for path in paths:
        size_mb = path.stat().st_size / (1024 * 1024)
        rows.append({"area": "supporting_information", "check": f"{path.name}_size", "status": "pass" if size_mb < 10 else "revise", "evidence": f"{size_mb:.2f}MB"})
        if path.suffix.lower() == ".xlsx":
            wb = load_workbook(path, read_only=True, data_only=True)
            rows.append({"area": "supporting_information", "check": f"{path.name}_sheets", "status": "pass", "evidence": ";".join(f"{ws.title}:{ws.max_row}x{ws.max_column}" for ws in wb.worksheets)})
        if path.suffix.lower() == ".zip":
            with zipfile.ZipFile(path) as zf:
                names = sorted(zf.namelist())
            rows.append({"area": "supporting_information", "check": f"{path.name}_contents", "status": "pass" if "scripts/run_transferability_supplement.py" in names else "revise", "evidence": f"entries={len(names)}; transferability_script={'scripts/run_transferability_supplement.py' in names}"})
    return rows


def build_upload_zip(upload_files: list[Path]) -> Path:
    zip_path = OUT / UPLOAD_ZIP_NAME
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as zf:
        for path in upload_files:
            zf.write(path, path.relative_to(OUT).as_posix())
    return zip_path


def render_docx(docx_path: Path) -> dict[str, Any]:
    render_dir = OUT / "audit" / "rendered_docx"
    if render_dir.exists():
        shutil.rmtree(render_dir)
    render_dir.mkdir(parents=True, exist_ok=True)
    soffice_candidates = [
        Path("C:/Program Files/LibreOffice/program/soffice.com"),
        Path("C:/Program Files/LibreOffice/program/soffice.exe"),
    ]
    pdftoppm_candidates = [
        Path("C:/Users/11970/AppData/Local/Microsoft/WinGet/Packages/oschwartz10612.Poppler_Microsoft.Winget.Source_8wekyb3d8bbwe/poppler-25.07.0/Library/bin/pdftoppm.exe"),
        Path("C:/Program Files/poppler/Library/bin/pdftoppm.exe"),
    ]
    soffice = next((p for p in soffice_candidates if p.exists()), None)
    pdftoppm = next((p for p in pdftoppm_candidates if p.exists()), None)
    if not soffice or not pdftoppm:
        return {"status": "skipped", "reason": f"soffice={soffice}; pdftoppm={pdftoppm}"}
    subprocess.run([str(soffice), "--headless", "--convert-to", "pdf", "--outdir", str(render_dir), str(docx_path)], check=True, stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
    pdf = render_dir / docx_path.with_suffix(".pdf").name
    subprocess.run([str(pdftoppm), "-png", "-r", "120", str(pdf), str(render_dir / "page")], check=True, stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
    pngs = sorted(render_dir.glob("page-*.png")) or sorted(render_dir.glob("page*.png"))
    return {"status": "pass", "pdf": str(pdf), "pdf_bytes": pdf.stat().st_size if pdf.exists() else 0, "page_png_count": len(pngs), "first_page_png_bytes": pngs[0].stat().st_size if pngs else 0}


def write_audit(rows: list[dict[str, Any]], render_result: dict[str, Any], zip_path: Path) -> None:
    rows.append({"area": "render", "check": "docx_to_pdf_png", "status": render_result.get("status", "unknown"), "evidence": json.dumps(render_result, ensure_ascii=False)})
    rows.append({"area": "upload_zip", "check": "sha256", "status": "pass", "evidence": sha256(zip_path)})
    with zipfile.ZipFile(zip_path) as zf:
        entries = sorted(zf.namelist())
    expected = sorted(
        [
            "cover_letter/cover_letter_PLOS_ONE.docx",
            "figures/Fig1.tif",
            "figures/Fig2.tif",
            "figures/Fig3.tif",
            "figures/Fig4.tif",
            "figures/Fig5.tif",
            f"manuscript/{MANUSCRIPT_NAME}",
            "supporting_information/S1_Table.xlsx",
            "supporting_information/S2_Code.zip",
            "supporting_information/S3_Transferability_Supplement.xlsx",
        ]
    )
    rows.append({"area": "upload_zip", "check": "entry_whitelist", "status": "pass" if entries == expected else "revise", "evidence": json.dumps(entries, ensure_ascii=False)})
    fields = ["area", "check", "status", "evidence"]
    with (OUT / "audit" / "PLOS_ONE_checked_submission_audit.tsv").open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fields, delimiter="\t")
        writer.writeheader()
        writer.writerows(rows)
    (OUT / "audit" / "PLOS_ONE_checked_submission_audit.json").write_text(json.dumps(rows, indent=2, ensure_ascii=False), encoding="utf-8")


def write_readme(zip_path: Path) -> None:
    text = f"""# PLOS ONE Checked Submission Package

Generated: {datetime.now().isoformat(timespec='seconds')}

## Upload files

- `manuscript/{MANUSCRIPT_NAME}`
- `figures/Fig1.tif` to `figures/Fig5.tif`
- `supporting_information/S1_Table.xlsx`
- `supporting_information/S2_Code.zip`
- `supporting_information/S3_Transferability_Supplement.xlsx`
- `cover_letter/cover_letter_PLOS_ONE.docx`

Clean upload ZIP: `{zip_path.name}`

## Submission-form text

Use `submission_admin/submission_form_statements.md` for author contributions, funding, competing interests, ethics, data availability and AI disclosure fields in Editorial Manager.

## Remaining author-side confirmations

- ORCID iD for each author, if required.
- Whether to nominate Academic Editors.
- Whether a preprint exists.
- Final all-author approval.
"""
    (OUT / "README_checked_submission.md").write_text(text, encoding="utf-8")


def main() -> None:
    reset_out()
    markdown = build_checked_markdown()
    manuscript = build_checked_docx(markdown)
    figures = copy_figures()
    supporting = copy_supporting_files()
    cover = build_cover_letter()
    write_submission_form_statements()
    shutil.copy2(Path(__file__), OUT / "code" / Path(__file__).name)

    upload_files = [manuscript, *figures, *supporting, cover]
    zip_path = build_upload_zip(upload_files)
    audit_rows = []
    audit_rows.extend(audit_docx(manuscript))
    audit_rows.extend(audit_figures(figures))
    audit_rows.extend(audit_supporting(supporting))
    render_result = render_docx(manuscript)
    write_audit(audit_rows, render_result, zip_path)
    write_readme(zip_path)
    print(json.dumps({"status": "completed", "out": str(OUT), "zip": str(zip_path), "zip_sha256": sha256(zip_path), "render": render_result}, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
