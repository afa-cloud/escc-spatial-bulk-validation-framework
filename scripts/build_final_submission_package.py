#!/usr/bin/env python
"""Build a journal-facing English submission package for the ESCC project.

Outputs:
- English manuscript in Markdown and DOCX.
- Main figures as PNG and SVG.
- Supplementary table workbook plus machine-readable TSV copies.
- Reference list in TSV and BibTeX.
- Cover letter and package README.

The package deliberately preserves reviewer claim limits: no independent
survival biomarker claim, no direct JAG1-NOTCH1 ligand-receptor claim, and no
drug-sensitivity prediction claim.
"""

from __future__ import annotations

import csv
import base64
import json
import math
import re
import shutil
import textwrap
import zipfile
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_ROOT = ROOT / "spatial_escc_workflow"
TABLE_ROOT = OUT_ROOT / "results" / "tables"
FIG_ROOT = OUT_ROOT / "results" / "figures"
REVIEW_ROOT = OUT_ROOT / "reviews"
DELIVERABLE_ROOT = OUT_ROOT / "deliverables"

RUN_DATE = datetime.now(UTC).date().isoformat()
PACKAGE_ROOT = OUT_ROOT / "submission_ready_2026-04-26"
MANUSCRIPT_DIR = PACKAGE_ROOT / "manuscript"
FIGURE_DIR = PACKAGE_ROOT / "figures"
SUPP_DIR = PACKAGE_ROOT / "supplementary_tables"
REF_DIR = PACKAGE_ROOT / "references"
REVIEW_DIR = PACKAGE_ROOT / "review_audit"
CODE_DIR = PACKAGE_ROOT / "code"
PLOS_DIR = PACKAGE_ROOT / "plos_one_submission"
PLOS_MANUSCRIPT_DIR = PLOS_DIR / "manuscript"
PLOS_FIGURE_DIR = PLOS_DIR / "figures"
PLOS_SUPPORTING_DIR = PLOS_DIR / "supporting_information"
PLOS_REVIEW_DIR = PLOS_DIR / "internal_review"

TITLE = (
    "Spatially informed public transcriptomic and source-table integration "
    "identifies a reproducible CAF/ECM stromal-remodeling phenotype in "
    "esophageal squamous cell carcinoma"
)
SHORT_TITLE = "Spatially informed ESCC stromal-remodeling phenotype"


REFERENCES = [
    {
        "id": "Sung2021",
        "authors": "Sung H, Ferlay J, Siegel RL, Laversanne M, Soerjomataram I, Jemal A, Bray F.",
        "title": "Global Cancer Statistics 2020: GLOBOCAN estimates of incidence and mortality worldwide for 36 cancers in 185 countries.",
        "journal": "CA Cancer J Clin",
        "year": "2021",
        "volume": "71",
        "pages": "209-249",
        "doi": "10.3322/caac.21660",
        "pmid": "33538338",
    },
    {
        "id": "TCGA2017",
        "authors": "Cancer Genome Atlas Research Network.",
        "title": "Integrated genomic characterization of oesophageal carcinoma.",
        "journal": "Nature",
        "year": "2017",
        "volume": "541",
        "pages": "169-175",
        "doi": "10.1038/nature20805",
        "pmid": "28052061",
    },
    {
        "id": "Liu2023Spatial",
        "authors": "Liu X, Zhao S, Wang K, Zhou L, Jiang M, Gao Y, et al.",
        "title": "Spatial transcriptomics analysis of esophageal squamous precancerous lesions and their progression to esophageal cancer.",
        "journal": "Nat Commun",
        "year": "2023",
        "volume": "14",
        "pages": "4779",
        "doi": "10.1038/s41467-023-40343-5",
        "pmid": "37553345",
    },
    {
        "id": "Liu2026SpatialOmics",
        "authors": "Liu Z, Zhou W, Li L, Song C, Yue M, Lv H, et al.",
        "title": "Spatial omics study reveals molecular-cellular dynamics of tumor ecosystem in esophageal squamous-cell carcinoma initiation and progression.",
        "journal": "Cell Rep Med",
        "year": "2026",
        "volume": "7",
        "pages": "102650",
        "doi": "10.1016/j.xcrm.2026.102650",
        "pmid": "41791392",
    },
    {
        "id": "Sawada2015",
        "authors": "Sawada G, Niida A, Hirata H, Komatsu H, Uchi R, Shimamura T, et al.",
        "title": "An integrative analysis to identify driver genes in esophageal squamous cell carcinoma.",
        "journal": "PLoS One",
        "year": "2015",
        "volume": "10",
        "pages": "e0139808",
        "doi": "10.1371/journal.pone.0139808",
        "pmid": "26465158",
    },
    {
        "id": "Li2014",
        "authors": "Li J, Chen Z, Tian L, Zhou C, He MY, Gao Y, et al.",
        "title": "LncRNA profile study reveals a three-lncRNA signature associated with the survival of patients with oesophageal squamous cell carcinoma.",
        "journal": "Gut",
        "year": "2014",
        "volume": "63",
        "pages": "1700-1710",
        "doi": "10.1136/gutjnl-2013-305806",
        "pmid": "24522499",
    },
    {
        "id": "Barrett2013",
        "authors": "Barrett T, Wilhite SE, Ledoux P, Evangelista C, Kim IF, Tomashevsky M, et al.",
        "title": "NCBI GEO: archive for functional genomics data sets--update.",
        "journal": "Nucleic Acids Res",
        "year": "2013",
        "volume": "41",
        "pages": "D991-D995",
        "doi": "10.1093/nar/gks1193",
        "pmid": "23193258",
    },
    {
        "id": "Goldman2020",
        "authors": "Goldman MJ, Craft B, Hastie M, Repecka K, McDade F, Kamath A, et al.",
        "title": "Visualizing and interpreting cancer genomics data via the Xena platform.",
        "journal": "Nat Biotechnol",
        "year": "2020",
        "volume": "38",
        "pages": "675-678",
        "doi": "10.1038/s41587-020-0546-8",
        "pmid": "32444850",
    },
    {
        "id": "Vivian2017",
        "authors": "Vivian J, Rao AA, Nothaft FA, Ketchum C, Armstrong J, Novak A, et al.",
        "title": "Toil enables reproducible, open source, big biomedical data analyses.",
        "journal": "Nat Biotechnol",
        "year": "2017",
        "volume": "35",
        "pages": "314-316",
        "doi": "10.1038/nbt.3772",
        "pmid": "28398314",
    },
    {
        "id": "Yates2015",
        "authors": "Yates A, Beal K, Keenan S, McLaren W, Pignatelli M, Ritchie GR, et al.",
        "title": "The Ensembl REST API: Ensembl data for any language.",
        "journal": "Bioinformatics",
        "year": "2015",
        "volume": "31",
        "pages": "143-145",
        "doi": "10.1093/bioinformatics/btu613",
        "pmid": "25236461",
    },
    {
        "id": "Iorio2016",
        "authors": "Iorio F, Knijnenburg TA, Vis DJ, Bignell GR, Menden MP, Schubert M, et al.",
        "title": "A landscape of pharmacogenomic interactions in cancer.",
        "journal": "Cell",
        "year": "2016",
        "volume": "166",
        "pages": "740-754",
        "doi": "10.1016/j.cell.2016.06.017",
        "pmid": "27397505",
    },
    {
        "id": "Yates2020",
        "authors": "Yates AD, Achuthan P, Akanni W, Allen J, Allen J, Alvarez-Jarreta J, et al.",
        "title": "Ensembl 2020.",
        "journal": "Nucleic Acids Res",
        "year": "2020",
        "volume": "48",
        "pages": "D682-D688",
        "doi": "10.1093/nar/gkz966",
        "pmid": "31691826",
    },
    {
        "id": "Hanahan2022",
        "authors": "Hanahan D.",
        "title": "Hallmarks of cancer: new dimensions.",
        "journal": "Cancer Discov",
        "year": "2022",
        "volume": "12",
        "pages": "31-46",
        "doi": "10.1158/2159-8290.CD-21-1059",
        "pmid": "35022204",
    },
    {
        "id": "Benjamini1995",
        "authors": "Benjamini Y, Hochberg Y.",
        "title": "Controlling the false discovery rate: a practical and powerful approach to multiple testing.",
        "journal": "J R Stat Soc Series B",
        "year": "1995",
        "volume": "57",
        "pages": "289-300",
        "doi": "10.1111/j.2517-6161.1995.tb02031.x",
        "pmid": "",
    },
    {
        "id": "REMARK2005",
        "authors": "McShane LM, Altman DG, Sauerbrei W, Taube SE, Gion M, Clark GM.",
        "title": "Reporting recommendations for tumor marker prognostic studies (REMARK).",
        "journal": "J Natl Cancer Inst",
        "year": "2005",
        "volume": "97",
        "pages": "1180-1184",
        "doi": "10.1093/jnci/dji237",
        "pmid": "16106022",
    },
]


def ensure_dirs() -> None:
    if PACKAGE_ROOT.exists():
        shutil.rmtree(PACKAGE_ROOT)
    for path in [
        MANUSCRIPT_DIR,
        FIGURE_DIR,
        SUPP_DIR,
        REF_DIR,
        REVIEW_DIR,
        CODE_DIR,
        PLOS_MANUSCRIPT_DIR,
        PLOS_FIGURE_DIR,
        PLOS_SUPPORTING_DIR,
        PLOS_REVIEW_DIR,
    ]:
        path.mkdir(parents=True, exist_ok=True)


def read_tsv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path, sep="\t")


def safe_float(value: Any) -> float:
    try:
        val = float(value)
    except (TypeError, ValueError):
        return float("nan")
    return val


def fmt(value: Any, digits: int = 3) -> str:
    val = safe_float(value)
    if math.isfinite(val):
        if abs(val) < 0.001 and val != 0:
            return f"{val:.2e}"
        return f"{val:.{digits}g}"
    return str(value)


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current: list[str] = []
    for word in words:
        trial = " ".join(current + [word])
        width = draw.textbbox((0, 0), trial, font=font)[2]
        if width <= max_width or not current:
            current.append(word)
        else:
            lines.append(" ".join(current))
            current = [word]
    if current:
        lines.append(" ".join(current))
    return lines


def draw_wrapped(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    font: ImageFont.ImageFont,
    max_width: int,
    fill: str = "#111111",
    line_gap: int = 4,
) -> int:
    x, y = xy
    lines = wrap_text(draw, text, font, max_width)
    line_height = draw.textbbox((0, 0), "Ag", font=font)[3] + line_gap
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height
    return y


def new_canvas(width: int = 2400, height: int = 1500) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (width, height), "white")
    return image, ImageDraw.Draw(image)


def save_svg(path: Path, width: int, height: int, body: str) -> None:
    path.write_text(
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">\n'
        '<rect width="100%" height="100%" fill="white"/>\n'
        f"{body}\n</svg>\n",
        encoding="utf-8",
    )


def save_svg_with_embedded_png(path: Path, width: int, height: int, png_path: Path) -> None:
    data = base64.b64encode(png_path.read_bytes()).decode("ascii")
    save_svg(
        path,
        width,
        height,
        f'<image href="data:image/png;base64,{data}" x="0" y="0" width="{width}" height="{height}"/>',
    )


def color_for_value(value: float, vmin: float = -1.0, vmax: float = 1.0) -> str:
    if not math.isfinite(value):
        return "#e8e8e8"
    t = max(0.0, min(1.0, (value - vmin) / (vmax - vmin)))
    if t < 0.5:
        u = t / 0.5
        r = int(54 + (245 - 54) * u)
        g = int(92 + (245 - 92) * u)
        b = int(169 + (245 - 169) * u)
    else:
        u = (t - 0.5) / 0.5
        r = int(245 + (178 - 245) * u)
        g = int(245 + (24 - 245) * u)
        b = int(245 + (43 - 245) * u)
    return f"#{r:02x}{g:02x}{b:02x}"


def save_png(image: Image.Image, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, dpi=(300, 300))


def make_figure1() -> dict[str, str]:
    image, draw = new_canvas()
    header_font = load_font(32, True)
    text_font = load_font(27)
    small_font = load_font(23)
    boxes = [
        (90, 230, 440, 500, "Spatially nominated axes", "OGT/PI3K/TLS; CAF-Epi/JAG1-NOTCH1 candidate niche"),
        (560, 230, 910, 500, "Bulk ESCC testing", "TCGA squamous subset; GSE47404 independent ESCC expression"),
        (1030, 230, 1380, 500, "Patient rescue", "GSE53625 probe-sequence rescue; paired tumor-normal and survival audit"),
        (1500, 230, 1850, 500, "Spatial source tables", "HRA003627 ROI progression; HRA008846 DEG, abundance, ligand-receptor tables"),
        (1970, 230, 2320, 500, "Review gates", "Executor/reviewer separation; claim downgrading before submission"),
    ]
    for idx, (x1, y1, x2, y2, head, text) in enumerate(boxes):
        fill = ["#eaf3ff", "#eff8ed", "#fff4e5", "#f4efff", "#f5f5f5"][idx]
        draw.rounded_rectangle((x1, y1, x2, y2), radius=22, fill=fill, outline="#333333", width=3)
        draw.text((x1 + 28, y1 + 28), head, font=header_font, fill="#111111")
        draw_wrapped(draw, (x1 + 28, y1 + 85), text, text_font, x2 - x1 - 56)
        if idx < len(boxes) - 1:
            draw.line((x2 + 25, 365, boxes[idx + 1][0] - 25, 365), fill="#555555", width=5)
            draw.polygon([(boxes[idx + 1][0] - 25, 365), (boxes[idx + 1][0] - 55, 347), (boxes[idx + 1][0] - 55, 383)], fill="#555555")
    y = 650
    bullets = [
        "Main reusable signal: CAF/ECM stromal-remodeling phenotype replicated in TCGA ESCC and GSE47404.",
        "Auxiliary signal: OGT/PI3K-associated hypoxia and epithelial progression context.",
        "Positive validation: GSE53625 paired tumor-normal upshift after target probe rescue.",
        "Negative gates retained: no independent survival support, no direct JAG1-NOTCH1 ligand-receptor row, no drug-response prediction model.",
    ]
    draw.text((90, y), "Final claim boundary", font=header_font, fill="#111111")
    y += 70
    for bullet in bullets:
        draw.ellipse((100, y + 8, 116, y + 24), fill="#2b7a78")
        y = draw_wrapped(draw, (135, y), bullet, text_font, 2000) + 18
    draw.text((90, 1370), "All results are association/source-table reproducibility evidence, not causal mechanism proof.", font=small_font, fill="#555555")
    png = FIGURE_DIR / "Figure1_workflow_evidence_gate.png"
    save_png(image, png)
    svg_body = "\n".join(
        [
            *[
                f'<rect x="{x1}" y="{y1}" width="{x2-x1}" height="{y2-y1}" rx="22" fill="{["#eaf3ff", "#eff8ed", "#fff4e5", "#f4efff", "#f5f5f5"][idx]}" stroke="#333" stroke-width="3"/>'
                f'<text x="{x1+28}" y="{y1+65}" font-size="32" font-family="Arial" font-weight="700">{head}</text>'
                f'<text x="{x1+28}" y="{y1+120}" font-size="24" font-family="Arial">{text}</text>'
                for idx, (x1, y1, x2, y2, head, text) in enumerate(boxes)
            ],
        ]
    )
    svg = FIGURE_DIR / "Figure1_workflow_evidence_gate.svg"
    save_svg(svg, 2400, 1500, svg_body)
    return {"png": str(png), "svg": str(svg)}


def association_lookup() -> pd.DataFrame:
    frames = [
        read_tsv(TABLE_ROOT / "deep_axis_tcga_immune_associations.tsv"),
        read_tsv(TABLE_ROOT / "deep_axis_tcga_pathway_associations.tsv"),
        read_tsv(TABLE_ROOT / "deep_axis_geo_gse47404_associations.tsv"),
    ]
    df = pd.concat(frames, ignore_index=True)
    return df


def make_figure2() -> dict[str, str]:
    df = association_lookup()
    rows = [
        ("TCGA_ESCC_Xena", "caf_epi_jag1_notch_niche", "TCGA CAF/ECM axis"),
        ("GSE47404", "caf_epi_jag1_notch_niche", "GSE47404 CAF/ECM axis"),
        ("TCGA_ESCC_Xena", "ogt_pi3k_tls_axis", "TCGA OGT/PI3K axis"),
        ("GSE47404", "ogt_pi3k_tls_axis", "GSE47404 OGT/PI3K axis"),
    ]
    panels = ["caf", "ecm_remodeling", "emt", "hypoxia", "pi3k_akt_mtor", "notch", "tls_b_cell", "macrophage_m2", "endothelial"]
    values: list[list[float]] = []
    fdrs: list[list[float]] = []
    for dataset, axis, _ in rows:
        row_vals: list[float] = []
        row_fdrs: list[float] = []
        for panel in panels:
            sub = df[(df["dataset"] == dataset) & (df["axis_id"] == axis) & (df["panel_id"] == panel)]
            if sub.empty:
                row_vals.append(float("nan"))
                row_fdrs.append(float("nan"))
            else:
                item = sub.iloc[0]
                row_vals.append(safe_float(item["spearman_rho"]))
                row_fdrs.append(safe_float(item["spearman_fdr_approx"]))
        values.append(row_vals)
        fdrs.append(row_fdrs)
    image, draw = new_canvas(2500, 1500)
    label_font = load_font(26, True)
    text_font = load_font(23)
    x0, y0 = 500, 220
    cell_w, cell_h = 190, 170
    for j, panel in enumerate(panels):
        label = panel.replace("_", " ")
        draw_wrapped(draw, (x0 + j * cell_w + 8, y0 - 100), label, text_font, cell_w - 16)
    for i, (_, _, label) in enumerate(rows):
        draw_wrapped(draw, (95, y0 + i * cell_h + 50), label, label_font, 360)
        for j, val in enumerate(values[i]):
            x = x0 + j * cell_w
            y = y0 + i * cell_h
            fill = color_for_value(val)
            draw.rectangle((x, y, x + cell_w - 6, y + cell_h - 6), fill=fill, outline="#ffffff")
            txt = "NA" if not math.isfinite(val) else f"{val:.2f}"
            draw.text((x + 55, y + 50), txt, font=label_font, fill="#111111")
            if math.isfinite(fdrs[i][j]) and fdrs[i][j] < 0.05:
                draw.text((x + 80, y + 95), "*", font=load_font(40, True), fill="#111111")
    legend_x, legend_y = 520, 1050
    draw.text((legend_x, legend_y), "Spearman rho", font=label_font, fill="#111111")
    for k in range(120):
        val = -1 + 2 * k / 119
        draw.rectangle((legend_x + 230 + 4 * k, legend_y + 5, legend_x + 230 + 4 * k + 4, legend_y + 45), fill=color_for_value(val))
    draw.text((legend_x + 230, legend_y + 55), "-1", font=text_font, fill="#333333")
    draw.text((legend_x + 455, legend_y + 55), "0", font=text_font, fill="#333333")
    draw.text((legend_x + 690, legend_y + 55), "+1", font=text_font, fill="#333333")
    draw.text((90, 1330), "* FDR < 0.05. The CAF/ECM signal is reproduced across TCGA ESCC and GSE47404; the OGT/PI3K axis is most consistently hypoxia-linked.", font=text_font, fill="#555555")
    png = FIGURE_DIR / "Figure2_bulk_axis_association_heatmap.png"
    save_png(image, png)
    svg_lines = []
    for i, (_, _, label) in enumerate(rows):
        svg_lines.append(f'<text x="95" y="{y0 + i * cell_h + 90}" font-size="28" font-family="Arial" font-weight="700">{label}</text>')
        for j, val in enumerate(values[i]):
            x = x0 + j * cell_w
            y = y0 + i * cell_h
            svg_lines.append(f'<rect x="{x}" y="{y}" width="{cell_w - 6}" height="{cell_h - 6}" fill="{color_for_value(val)}" stroke="#fff"/>')
            svg_lines.append(f'<text x="{x+58}" y="{y+82}" font-size="30" font-family="Arial">{fmt(val,2)}</text>')
    svg = FIGURE_DIR / "Figure2_bulk_axis_association_heatmap.svg"
    save_svg(svg, 2500, 1500, "\n".join(svg_lines))
    return {"png": str(png), "svg": str(svg)}


def make_figure3() -> dict[str, str]:
    paired = read_tsv(TABLE_ROOT / "gse53625_rescue_tumor_normal_validation.tsv")
    surv = read_tsv(TABLE_ROOT / "gse53625_rescue_survival_validation.tsv")
    image, draw = new_canvas(2400, 1500)
    label_font = load_font(30, True)
    text_font = load_font(25)
    small_font = load_font(22)
    gene_font = load_font(18)
    axes = [
        ("ogt_pi3k_tls_axis", "OGT/PI3K"),
        ("caf_epi_jag1_notch_niche", "CAF/ECM"),
    ]
    # Panel A: accepted probes.
    draw.text((90, 105), "A. Probe rescue coverage", font=label_font, fill="#111111")
    cov = read_tsv(TABLE_ROOT / "gse53625_rescue_axis_gene_coverage.tsv")
    x_base, y_base = 120, 520
    max_count = max(cov["accepted_probe_count"].astype(float).max(), 1)
    genes = ["OGT", "PIK3CA", "CCND1", "LAMB1", "JAG1", "NOTCH1", "FAP", "COL1A1", "COL1A2", "POSTN", "CXCL8"]
    for idx, gene in enumerate(genes):
        sub = cov[cov["gene_symbol"] == gene].iloc[0]
        count = safe_float(sub["accepted_probe_count"])
        bar_h = int(260 * count / max_count)
        x = x_base + idx * 92
        draw.rectangle((x, y_base - bar_h, x + 50, y_base), fill="#5b8def")
        draw.text((x - 5, y_base + 14), gene, font=gene_font, fill="#333333")
        draw.text((x + 8, y_base - bar_h - 32), str(int(count)), font=small_font, fill="#333333")
    draw_wrapped(draw, (130, 590), "All axis genes had at least two accepted probes; rescue is target-scope, not genome-wide specificity audit.", small_font, 900, fill="#555555")
    # Panel B: tumor-normal deltas.
    draw.text((90, 740), "B. Paired tumor-normal score shifts", font=label_font, fill="#111111")
    for idx, (axis_id, label) in enumerate(axes):
        row = paired[paired["axis_id"] == axis_id].iloc[0]
        delta = safe_float(row["median_tumor_minus_normal_axis_score"])
        pval = safe_float(row["two_sided_sign_test_p"])
        x = 160 + idx * 420
        y0 = 1250
        bar_h = int(250 * delta / 1.6)
        draw.rectangle((x, y0 - bar_h, x + 180, y0), fill="#2b7a78")
        draw.text((x, y0 + 25), label, font=label_font, fill="#111111")
        draw.text((x - 10, y0 - bar_h - 70), f"median Δ={delta:.2f}", font=text_font, fill="#111111")
        draw.text((x - 10, y0 - bar_h - 35), f"p={pval:.1e}", font=small_font, fill="#333333")
    draw.line((120, 1250, 880, 1250), fill="#333333", width=3)
    # Panel C: survival p-values.
    draw.text((1250, 105), "C. Survival gate", font=label_font, fill="#111111")
    for idx, (axis_id, label) in enumerate(axes):
        row = surv[surv["axis_id"] == axis_id].iloc[0]
        pval = safe_float(row["logrank_p"])
        rr = safe_float(row["event_rate_ratio_approx"])
        y = 230 + idx * 260
        draw.rounded_rectangle((1250, y, 2240, y + 190), radius=24, fill="#fff4e5", outline="#b75d00", width=3)
        draw.text((1285, y + 30), label, font=label_font, fill="#111111")
        draw.text((1285, y + 85), f"log-rank p={pval:.3f}; event-rate ratio={rr:.2f}", font=text_font, fill="#111111")
        draw.text((1285, y + 125), "Status: no survival support", font=text_font, fill="#b00020")
    draw_wrapped(draw, (1250, 790), "Reviewer decision: use GSE53625 as paired tumor-normal validation only; do not claim independent prognosis.", text_font, 920, fill="#555555")
    png = FIGURE_DIR / "Figure3_gse53625_probe_rescue_validation.png"
    save_png(image, png)
    svg = FIGURE_DIR / "Figure3_gse53625_probe_rescue_validation.svg"
    save_svg_with_embedded_png(svg, 2400, 1500, png)
    return {"png": str(png), "svg": str(svg)}


def make_figure4() -> dict[str, str]:
    h3 = read_tsv(TABLE_ROOT / "hra003627_source_table_quantification.tsv")
    deg = read_tsv(TABLE_ROOT / "hra008846_deg_axis_hits.tsv")
    lr = read_tsv(TABLE_ROOT / "hra008846_ligand_receptor_axis_hits.tsv")
    image, draw = new_canvas(2500, 1600)
    label_font = load_font(30, True)
    text_font = load_font(24)
    small_font = load_font(21)
    # Panel A line plot.
    draw.text((90, 175), "A. HRA003627 ROI-level progression signatures", font=label_font, fill="#111111")
    plot_x, plot_y, plot_w, plot_h = 120, 310, 820, 470
    draw.rectangle((plot_x, plot_y, plot_x + plot_w, plot_y + plot_h), outline="#333333", width=3)
    stages = ["Normal", "LGIN", "HGIN", "ESCC"]
    colors = {"dk_keratinization": "#3366aa", "cancerization_progression": "#bb3333"}
    for sig, color in colors.items():
        sub = h3[h3["signature_id"] == sig]
        pts = []
        vals = []
        for stage in stages:
            row = sub[sub["stage"] == stage].iloc[0]
            vals.append(safe_float(row["mean_signature_z"]))
        vmin, vmax = -1.0, 1.3
        for idx, value in enumerate(vals):
            x = plot_x + 90 + idx * 210
            y = plot_y + plot_h - int((value - vmin) / (vmax - vmin) * (plot_h - 70)) - 35
            pts.append((x, y))
        for a, b in zip(pts, pts[1:]):
            draw.line((a[0], a[1], b[0], b[1]), fill=color, width=6)
        for x, y in pts:
            draw.ellipse((x - 10, y - 10, x + 10, y + 10), fill=color)
        draw.text((plot_x + 520, plot_y + (20 if sig == "dk_keratinization" else 60)), sig.replace("_", " "), font=small_font, fill=color)
    for idx, stage in enumerate(stages):
        draw.text((plot_x + 55 + idx * 210, plot_y + plot_h + 20), stage, font=small_font, fill="#333333")
    draw.text((plot_x + 15, plot_y + 18), "mean z", font=small_font, fill="#333333")
    # Panel B top DEG.
    draw.text((1120, 175), "B. HRA008846 axis-relevant source-table rows", font=label_font, fill="#111111")
    wanted = ["COL1A1", "POSTN", "COL1A2", "FAP", "OGT", "CCND1", "LAMB1", "SPP1"]
    top_rows = []
    for gene in wanted:
        sub = deg[deg["gene_symbol"] == gene].copy()
        if sub.empty:
            continue
        sub["_abs"] = sub["logFC"].apply(lambda x: abs(safe_float(x)))
        top_rows.append(sub.sort_values("_abs", ascending=False).iloc[0])
    x0, y0 = 1160, 780
    for idx, row in enumerate(top_rows):
        val = safe_float(row["logFC"])
        x_mid = 1660
        y = y0 - idx * 70
        bar_w = int(260 * min(abs(val), 3.5) / 3.5)
        fill = "#bb3333" if val > 0 else "#3366aa"
        if val >= 0:
            draw.rectangle((x_mid, y - 20, x_mid + bar_w, y + 20), fill=fill)
        else:
            draw.rectangle((x_mid - bar_w, y - 20, x_mid, y + 20), fill=fill)
        draw.text((1120, y - 20), row["gene_symbol"], font=text_font, fill="#111111")
        draw.text((1940, y - 20), f"logFC={val:.2f}", font=small_font, fill="#333333")
    draw.line((1660, 210, 1660, 820), fill="#333333", width=2)
    draw_wrapped(draw, (1120, 880), "Rows include DEG and shOGT perturbation source-table entries; p/FDR status is reported in Supplementary Table S9.", small_font, 1100, fill="#555555")
    # Panel C LR boundary.
    draw.text((90, 1000), "C. Ligand-receptor boundary from HRA008846 Table S6", font=label_font, fill="#111111")
    total_lr = len(lr)
    direct = int((lr["direct_jag1_notch1_flag"].astype(str) == "yes").sum())
    draw.rounded_rectangle((130, 1120, 1050, 1390), radius=30, fill="#f5f5f5", outline="#333333", width=3)
    draw.text((180, 1160), f"Axis-relevant ligand-receptor rows: {total_lr}", font=label_font, fill="#111111")
    draw.text((180, 1225), f"Direct JAG1-NOTCH1 rows: {direct}", font=label_font, fill="#b00020")
    draw_wrapped(draw, (180, 1290), "Interpretation: HRA008846 supports SPP1/POSTN/CXCL13-related microenvironment context, but not direct CAF-epithelial JAG1-NOTCH1 signaling.", text_font, 800, fill="#333333")
    png = FIGURE_DIR / "Figure4_spatial_source_table_evidence.png"
    save_png(image, png)
    svg = FIGURE_DIR / "Figure4_spatial_source_table_evidence.svg"
    save_svg_with_embedded_png(svg, 2500, 1600, png)
    return {"png": str(png), "svg": str(svg)}


def make_figure5() -> dict[str, str]:
    review = read_tsv(REVIEW_ROOT / "independent_patient_and_spatial_quant_review.tsv")
    image, draw = new_canvas(2400, 1500)
    label_font = load_font(31, True)
    text_font = load_font(24)
    small_font = load_font(21)
    rows = [
        ("CAF/ECM remodeling reproducibility", "PASS", "TCGA and GSE47404 correlations; HRA008846 stromal DEG rows"),
        ("OGT/PI3K-hypoxia auxiliary axis", "PASS WITH LIMITS", "Hypoxia replicated; PI3K/TLS specificity weaker"),
        ("GSE53625 paired tumor-normal support", "PASS WITH LIMITS", "179 pairs; both axes elevated in tumors"),
        ("Independent survival prediction", "NO SUPPORT", "GSE53625 log-rank p=0.258 and 0.373"),
        ("Direct JAG1-NOTCH1 ligand-receptor mechanism", "NO SUPPORT", "HRA008846 Table S6 direct rows=0"),
        ("Drug sensitivity prediction", "NO SUPPORT", "GDSC2 target-class context only"),
    ]
    x1, y = 120, 210
    for label, status, reason in rows:
        status_color = {
            "PASS": "#2b7a78",
            "PASS WITH LIMITS": "#b75d00",
            "NO SUPPORT": "#b00020",
        }[status]
        draw.rounded_rectangle((x1, y, 2280, y + 145), radius=24, fill="#ffffff", outline="#cccccc", width=3)
        draw.text((x1 + 35, y + 35), label, font=label_font, fill="#111111")
        draw.rounded_rectangle((x1 + 980, y + 35, x1 + 1400, y + 95), radius=18, fill=status_color)
        draw.text((x1 + 1005, y + 52), status, font=small_font, fill="white")
        draw_wrapped(draw, (x1 + 1450, y + 30), reason, text_font, 760, fill="#333333")
        y += 175
    draw.text((120, 1340), "Submission label: association and source-table reproducibility article; revise before submission, not a biomarker/mechanism claim.", font=text_font, fill="#555555")
    png = FIGURE_DIR / "Figure5_claim_boundary_audit.png"
    save_png(image, png)
    svg = FIGURE_DIR / "Figure5_claim_boundary_audit.svg"
    save_svg_with_embedded_png(svg, 2400, 1500, png)
    return {"png": str(png), "svg": str(svg)}


def make_all_figures() -> list[dict[str, str]]:
    figure_entries = []
    for idx, maker in enumerate([make_figure1, make_figure2, make_figure3, make_figure4, make_figure5], start=1):
        paths = maker()
        figure_entries.append({"figure": f"Figure {idx}", **paths})
    return figure_entries


def references_markdown() -> str:
    lines = []
    for idx, ref in enumerate(REFERENCES, start=1):
        doi = f" doi:{ref['doi']}." if ref.get("doi") else ""
        pmid = f" PMID:{ref['pmid']}." if ref.get("pmid") else ""
        lines.append(
            f"{idx}. {ref['authors']} {ref['title']} {ref['journal']}. "
            f"{ref['year']};{ref['volume']}:{ref['pages']}.{doi}{pmid}"
        )
    return "\n".join(lines)


def cite(ref_id: str) -> str:
    for idx, ref in enumerate(REFERENCES, start=1):
        if ref["id"] == ref_id:
            return f"[{idx}]"
    raise KeyError(ref_id)


def dataset_summary_table() -> pd.DataFrame:
    return pd.DataFrame(
        [
            {
                "Dataset/resource": "TCGA ESCC via UCSC Xena/Toil",
                "Sample scope": "92 squamous TCGA ESCA primary tumors",
                "Role": "Discovery and primary bulk association testing",
                "Claim boundary": "Associations only; no causality",
            },
            {
                "Dataset/resource": "GSE47404",
                "Sample scope": "71 ESCC tumors",
                "Role": "Independent bulk expression replication",
                "Claim boundary": "Tumor-only; no survival endpoint in matrix metadata",
            },
            {
                "Dataset/resource": "GSE53625",
                "Sample scope": "179 paired ESCC tumor-normal samples",
                "Role": "Target probe rescue, paired tumor-normal and survival audit",
                "Claim boundary": "Tumor-normal support only; survival not supported",
            },
            {
                "Dataset/resource": "HRA003627 source table",
                "Sample scope": "42 ROI-level rows across normal, low-grade, high-grade, cancer",
                "Role": "Spatial progression source-table quantification",
                "Claim boundary": "Source-table quantification; not raw spatial reanalysis",
            },
            {
                "Dataset/resource": "HRA008846 supplementary tables",
                "Sample scope": "Published DEG, abundance and ligand-receptor source tables",
                "Role": "Spatial source-table validation of OGT/CCND1 and CAF/ECM context",
                "Claim boundary": "No direct JAG1-NOTCH1 LR evidence",
            },
            {
                "Dataset/resource": "GDSC2 ESCA cell lines",
                "Sample scope": "35 ESCA cell lines in fitted dose-response table",
                "Role": "Drug target-class context",
                "Claim boundary": "No axis-expression drug-sensitivity prediction model",
            },
        ]
    )


def evidence_gate_table() -> pd.DataFrame:
    return pd.DataFrame(
        [
            {
                "Claim": "CAF/ECM stromal-remodeling phenotype is reproducible",
                "Evidence": "TCGA CAF rho=0.906/FDR=1.02e-16; TCGA ECM rho=0.917/FDR=3.80e-17; GSE47404 ECM rho=0.829/FDR=7.31e-11",
                "Status": "Pass",
            },
            {
                "Claim": "OGT/PI3K axis is hypoxia-linked",
                "Evidence": "TCGA hypoxia rho=0.651/FDR=3.14e-09; GSE47404 hypoxia rho=0.536/FDR=4.46e-05",
                "Status": "Pass with limits",
            },
            {
                "Claim": "Axes are elevated in ESCC tumor tissue",
                "Evidence": "GSE53625 paired tumor-normal median delta 0.823 and 1.482; sign-test p=6.10e-41 and 4.70e-52",
                "Status": "Pass with limits",
            },
            {
                "Claim": "Axes independently predict ESCC survival",
                "Evidence": "GSE53625 log-rank p=0.25785 and 0.372619",
                "Status": "No support",
            },
            {
                "Claim": "Direct CAF-epithelial JAG1-NOTCH1 signaling is proven",
                "Evidence": "HRA008846 Table S6 direct JAG1-NOTCH1 rows=0",
                "Status": "No support",
            },
            {
                "Claim": "Axis predicts drug sensitivity",
                "Evidence": "GDSC2 target-class summaries only; no joint expression-response model",
                "Status": "No support",
            },
        ]
    )


def manuscript_text() -> str:
    refs = references_markdown()
    return f"""# {TITLE}

Running title: {SHORT_TITLE}

Authors: to be added by the study team

Correspondence: to be added by the study team

## Abstract

### Background

Spatial transcriptomic studies have highlighted epithelial, stromal and immune-context programs during esophageal squamous cell carcinoma (ESCC) initiation and progression, but many spatial observations remain difficult to translate into reusable public-cohort evidence. We asked whether spatially nominated ESCC axes could be reproduced in public bulk transcriptomic cohorts and openly available spatial source tables without overstating prognostic or mechanistic claims.

### Methods

We evaluated two pre-specified axes: an OGT/PI3K/TLS axis and a CAF-epithelial JAG1-NOTCH1 candidate niche axis. TCGA ESCA samples were restricted to squamous histology and analyzed through UCSC Xena/Toil expression data. GSE47404 was used as an independent ESCC tumor-only expression cohort. GSE53625 was re-examined by target-scope Agilent probe-sequence rescue against Ensembl transcript and genomic sequences, enabling paired tumor-normal and survival audits. Open source tables from HRA003627 and HRA008846 were quantified to evaluate spatial progression signatures, cell abundance trends, differential expression rows and ligand-receptor evidence. Executor and reviewer gates were used to enforce claim boundaries.

### Results

The most reproducible signal was a CAF/ECM stromal-remodeling phenotype derived from the original CAF-epithelial/JAG1-NOTCH1 candidate axis. In TCGA ESCC, this score correlated strongly with CAF markers (Spearman rho=0.906, FDR=1.02e-16), ECM remodeling (rho=0.917, FDR=3.80e-17) and EMT (rho=0.818, FDR=5.46e-14). The same pattern replicated in GSE47404 for ECM remodeling (rho=0.829, FDR=7.31e-11), CAF markers (rho=0.795, FDR=5.14e-10) and EMT (rho=0.754, FDR=2.57e-09). The OGT/PI3K/TLS axis was most consistently linked to hypoxia in TCGA (rho=0.651, FDR=3.14e-09) and GSE47404 (rho=0.536, FDR=4.46e-05), whereas PI3K/TLS specificity was weaker. GSE53625 probe rescue accepted 98 target-scope probes covering 16 axis genes. Both axes were elevated in 179 paired tumors relative to adjacent normal tissue, but neither supported survival stratification. HRA003627 source tables supported loss of keratinization/differentiation and gain of cancerization/progression programs. HRA008846 source tables supported OGT, CCND1, LAMB1 and CAF/ECM-related spatial progression context, but direct JAG1-NOTCH1 ligand-receptor rows were absent.

### Conclusions

Spatially informed public-cohort integration supports a reproducible CAF/ECM stromal-remodeling phenotype in ESCC and an auxiliary OGT/PI3K-associated hypoxia/epithelial progression state. The present evidence supports association and source-table reproducibility claims, but not independent survival prediction, direct JAG1-NOTCH1 signaling or drug-sensitivity prediction.

Keywords: esophageal squamous cell carcinoma; spatial transcriptomics; CAF; ECM remodeling; OGT; PI3K; hypoxia; public transcriptomics; source table reproducibility

## Introduction

Esophageal cancer remains a major global cancer burden, and ESCC is a dominant histological subtype in many high-incidence regions {cite("Sung2021")}. Large-scale genomic studies have established that esophageal carcinoma comprises molecularly distinct histological entities, making histology-aware analyses essential when reusing public ESCA data {cite("TCGA2017")}. At the same time, ESCC progression reflects not only tumor-intrinsic alterations but also stromal remodeling, extracellular matrix (ECM) changes, immune contexture and hypoxia-linked epithelial stress, all of which are central themes in current cancer biology {cite("Hanahan2022")}.

Spatial transcriptomic studies have begun to map these processes in tissue context. HRA003627 profiled esophageal squamous precancerous lesions and their progression to carcinoma, providing ROI-level source data that can be reused to quantify progression-associated programs {cite("Liu2023Spatial")}. HRA008846 further reported molecular-cellular dynamics of the ESCC ecosystem during initiation and progression, including differential expression, cell abundance and ligand-receptor supplementary tables {cite("Liu2026SpatialOmics")}. However, raw spatial matrices are not always openly available, and many spatial observations are difficult to validate directly in patient-scale cohorts.

A conservative public-data strategy is therefore to convert spatially nominated observations into pre-specified gene-set scores, test them in public bulk ESCC cohorts, and then return to open spatial source tables for direct source-table reproducibility. This approach is not equivalent to raw spatial reanalysis, but it can help distinguish reproducible tissue-level phenotypes from mechanistic hypotheses requiring further validation. Here, we evaluate two spatially informed ESCC axes: an OGT/PI3K/TLS axis and a CAF-epithelial/JAG1-NOTCH1 candidate niche axis. We explicitly ask whether either axis has reproducible bulk, patient-paired and source-table support, and we use independent review gates to prevent overclaiming survival, drug-response or direct ligand-receptor mechanisms.

## Materials and methods

### Study design and claim-control framework

This was a public-data integrative bioinformatics study. No controlled-access raw sequencing data, wet-lab experiments or newly collected patient material were used. Analyses were designed to support association and source-table reproducibility claims only. Executor and reviewer identifiers were recorded in the project artifacts, and final claims were downgraded whenever evidence failed a pre-specified gate. A pre-submission evidence gate restricted the manuscript to association and source-table reproducibility claims, not prognostic biomarker, causal mechanism or therapeutic-response claims.

### Axis definitions

The OGT/PI3K/TLS axis consisted of OGT, PIK3CA, AKT1, CCND1, LAMB1, SPP1, KRT17, APOBEC3A, JAG1 and NOTCH1. The CAF-epithelial/JAG1-NOTCH1 candidate axis consisted of JAG1, NOTCH1, FAP, COL1A1, COL1A2, POSTN, CXCL1, CXCL8 and SPP1. During review, the second axis was reframed as a CAF/ECM stromal-remodeling phenotype because no direct JAG1-NOTCH1 ligand-receptor rows were found in the HRA008846 Table S6 source table.

### TCGA ESCC analysis

TCGA ESCA primary tumors were filtered to squamous histology using GDC-derived case metadata. Gene expression values were obtained from UCSC Xena/Toil resources {cite("Goldman2020")}{cite("Vivian2017")}. Axis scores were calculated as the mean expression of available axis genes. Immune and pathway panels included CAF, endothelial, macrophage M2, TLS/B-cell, hypoxia, PI3K/AKT/mTOR, Notch, EMT, ECM remodeling and proliferation programs. Associations were tested by Spearman correlation and by median-split Mann-Whitney comparisons. P values were adjusted by the Benjamini-Hochberg false discovery rate procedure {cite("Benjamini1995")}.

### Independent bulk validation in GSE47404

GSE47404, originally associated with an integrative ESCC driver-gene study {cite("Sawada2015")}, was downloaded from GEO {cite("Barrett2013")}. GPL6480 probe annotations were used to map probes to gene symbols, and multiple probes for the same gene were averaged. GSE47404 was treated as a tumor-only independent expression cohort suitable for correlation replication and limited pathology-associated checks, but not for tumor-normal or survival validation.

### GSE53625 target probe rescue

GSE53625 was downloaded from GEO as a 358-sample paired ESCC expression resource linked to the original lncRNA survival-signature study {cite("Li2014")}. The public series matrix uses Agilent feature numbers. We therefore used the first RAW member to extract probe sequences and performed target-scope rescue against Ensembl transcript cDNA and genomic sequences obtained through the Ensembl REST API {cite("Yates2015")}{cite("Yates2020")}. Exact matches and one-mismatch matches were accepted only when uniquely assigned within the pre-specified axis gene universe. This procedure is a target-scope rescue and not a genome-wide probe-specificity audit.

Paired tumor-normal differences were summarized by median tumor-minus-normal axis scores and two-sided sign tests. Survival audits used tumor samples only, median score cutpoints, log-rank tests and approximate event-rate ratios. Survival support required log-rank P<0.05.

### Spatial source-table quantification

For HRA003627, ROI-level source data were used to quantify a keratinization/differentiation signature (CRNN and MAL) and a cancerization/progression signature (TAGLN2, KRT16, KRT17, S100A8, TOP2A, MKI67, LAMC2, CCN2, ANO1, ITGA6 and MMP14). Stage trends across normal, low-grade, high-grade and cancer ROIs were tested by Spearman correlation.

For HRA008846, Table S3 was parsed for axis-relevant differential-expression or perturbation rows, Table S4 for cell abundance trends and Table S6 for ligand-receptor rows. Rows involving SPP1, POSTN, CXCL13/CXCR5 and other axis-context genes were retained, and a specific flag was used to search for direct JAG1-NOTCH1 ligand-receptor evidence. An initial automated PMC supplementary-file download returned invalid HTML responses; the valid cached XLSX files ultimately used for Table S3, Table S4 and Table S6 are documented in the data-source manifest.

### Drug-response context

GDSC2 fitted dose-response data were filtered to ESCA cell lines and summarized for axis-relevant drug target classes {cite("Iorio2016")}. These data were interpreted as target-class context only, because no axis-expression and drug-response joint model was fitted.

## Results

### A CAF/ECM stromal-remodeling phenotype is the dominant reproducible signal

The original CAF-epithelial/JAG1-NOTCH1 candidate axis showed its strongest bulk signal as a stromal remodeling phenotype rather than as direct ligand-receptor evidence. In TCGA ESCC, the score correlated with CAF markers (rho=0.906, FDR=1.02e-16), ECM remodeling (rho=0.917, FDR=3.80e-17) and EMT (rho=0.818, FDR=5.46e-14). The same pattern replicated in GSE47404, where the score correlated with ECM remodeling (rho=0.829, FDR=7.31e-11), CAF markers (rho=0.795, FDR=5.14e-10) and EMT (rho=0.754, FDR=2.57e-09). These cross-cohort associations nominate a reusable CAF/ECM stromal-remodeling phenotype in ESCC (Figure 2).

### The OGT/PI3K/TLS axis is most consistently hypoxia-linked

The OGT/PI3K/TLS axis was reproducibly associated with hypoxia. In TCGA ESCC, the OGT/PI3K/TLS score correlated with the hypoxia panel (rho=0.651, FDR=3.14e-09), and this association replicated in GSE47404 (rho=0.536, FDR=4.46e-05). By contrast, PI3K/AKT/mTOR, TLS/B-cell and Notch associations were weaker or less stable after adjustment. We therefore interpret this axis as an OGT/PI3K-associated hypoxia and epithelial progression context rather than as a specific TLS biomarker.

### GSE53625 probe rescue supports tumor-normal upshift but not survival prediction

Target probe rescue in GSE53625 accepted 98 probe rows and covered 16 axis-related genes. Both axes had complete gene coverage for their predefined gene lists. In 179 paired tumor-normal samples, the OGT/PI3K axis was elevated in tumor tissue with a median tumor-normal delta of 0.823 (two-sided sign-test P=6.10e-41). The CAF/ECM phenotype was also elevated in tumor tissue with a median delta of 1.482 (P=4.70e-52). However, survival audits were negative: log-rank P values were 0.25785 for the OGT/PI3K axis and 0.372619 for the CAF/ECM phenotype. Thus, GSE53625 supports tumor-normal expression shifts but not independent survival prediction (Figure 3).

### Spatial source tables support progression and stromal context but not direct JAG1-NOTCH1 signaling

HRA003627 source-table quantification showed a decline in the keratinization/differentiation signature across disease stage (rho=-0.776, approximate P=6.84e-07) and an increase in the cancerization/progression signature (rho=0.811, approximate P=2.09e-07). These results support spatial progression context but do not directly test the CAF/JAG1-NOTCH1 axis.

HRA008846 Table S3 contained 32 axis-relevant source-table rows, including OGT upregulation in advanced versus early epithelial regions (logFC=1.022, FDR=0.00111), CCND1 upregulation in advanced epithelial regions (logFC=2.044, FDR=0.000355), and CCND1 downregulation after shOGT perturbation in KYSE30 cells (replicate Q values 0.000782 and 3.78e-06). Stromal/CAF-context rows included COL1A1 (logFC=3.410, FDR=7.57e-06), POSTN (logFC=2.583, FDR=4.22e-05), COL1A2 (logFC=2.449, FDR=0.00173) and FAP (logFC=1.018, FDR=0.0126). Table S6 contained 60 axis-relevant ligand-receptor rows, mainly involving SPP1, POSTN and CXCL13/CXCR5 context, but no direct JAG1-NOTCH1 rows. Therefore, HRA008846 supports stromal remodeling and epithelial progression context but not a proven JAG1-NOTCH1 cell-cell communication mechanism (Figure 4).

### Drug-response summaries are target-class context only

GDSC2 ESCA summaries highlighted EGFR/ERBB inhibitors and PI3K/AKT/mTOR-related compounds among axis-relevant target classes. Sapitinib, gefitinib, lapatinib, erlotinib and osimertinib appeared among the leading EGFR/ERBB-context rows. Because no joint model connected axis-expression scores with drug-response metrics, these results should be interpreted only as therapeutic target-class context, not drug-sensitivity prediction.

## Discussion

This study supports a conservative, source-verifiable ESCC story: a spatially informed CAF/ECM stromal-remodeling phenotype is reproducible across TCGA ESCC, GSE47404 and HRA008846 source tables, while an OGT/PI3K-related axis marks a hypoxia and epithelial progression context. The strongest manuscript lead is not direct JAG1-NOTCH1 signaling but a CAF/ECM-rich remodeling phenotype. This distinction is important because bulk expression associations and source-table ligand-receptor summaries cannot establish causal cell-cell communication.

The GSE53625 rescue analysis adds a useful patient-paired layer. The initial obstacle was a lack of direct public feature-to-gene annotation for the Agilent feature-number matrix. Target-scope sequence rescue recovered axis gene coverage and showed that both axes are elevated in tumors relative to paired adjacent normal tissue. However, survival analyses were negative. This prevents a prognostic biomarker claim and aligns with REMARK principles that tumor marker and prognostic studies require careful validation and reporting {cite("REMARK2005")}.

The HRA003627 and HRA008846 analyses also clarify what can and cannot be claimed from public spatial resources. HRA003627 source tables support disease-stage shifts in keratinization and progression signatures. HRA008846 source tables support OGT/CCND1 and CAF/ECM source-table findings. Yet direct JAG1-NOTCH1 ligand-receptor rows were absent, and the publicly available tables are not a substitute for raw spatial matrix reanalysis.

Several limitations should be stated plainly. First, TCGA and GEO bulk data cannot resolve cell origin; stromal abundance, tumor purity and biological state changes may all contribute to CAF/ECM signal intensity. Second, GSE53625 probe rescue was restricted to pre-specified target genes and did not perform a genome-wide specificity audit. Third, HRA003627 and HRA008846 source-table analyses are secondary quantifications of published tables rather than reprocessing of raw spatial data. Fourth, the GDSC2 analysis did not model axis scores against drug response. Finally, the current evidence does not support survival prediction, direct JAG1-NOTCH1 signaling or therapeutic-response prediction.

In summary, the present package is suitable for a public transcriptomic and source-table reproducibility manuscript. The most defensible title-level claim is that ESCC contains a reproducible CAF/ECM stromal-remodeling phenotype supported by bulk and source-table evidence, with an auxiliary OGT/PI3K-hypoxia context. Stronger mechanistic or clinical biomarker claims require independent clinical cohorts, raw spatial reanalysis or experimental validation.

## Data availability

All data used in this study are publicly available from TCGA/GDC/UCSC Xena, GEO accessions GSE47404 and GSE53625, HRA003627 source data, HRA008846 supplementary tables and GDSC2. Processed tables generated in this project are provided in the supplementary workbook and in machine-readable TSV files. Accession identifiers, URLs, download/cache status, local file sizes and source-table conflict notes are provided in the data-source manifest.

## Code availability

Reproducible scripts are included in the submission package under the `code/` directory. The main scripts are `run_spatial_axis_deep_validation.py`, `run_independent_patient_and_spatial_quant.py` and `build_final_submission_package.py`.

## Ethics statement

This study used only publicly available de-identified data and did not involve new human-subject recruitment, intervention or access to controlled raw sequencing data.

## Author contributions

Author contributions should be completed by the study team before journal submission.

## Funding

Funding information should be completed by the study team before journal submission.

## Competing interests

The authors should declare any competing interests before journal submission.

## Figure legends

**Figure 1. Evidence-gated integrative workflow.** Spatially nominated axes were evaluated in bulk ESCC cohorts, a target-rescued paired patient cohort, open spatial source tables and review gates. The final claim boundary limits the manuscript to association and source-table reproducibility.

**Figure 2. Cross-cohort axis associations with immune and pathway programs.** Spearman correlations show that the CAF/ECM phenotype is strongly associated with CAF, ECM remodeling and EMT programs in TCGA ESCC and GSE47404. The OGT/PI3K axis is most consistently associated with hypoxia. Asterisks indicate FDR<0.05.

**Figure 3. GSE53625 target probe rescue supports tumor-normal upshift but not survival prediction.** Target-scope probe rescue recovered axis gene coverage. Both axes were elevated in paired tumors relative to adjacent normal tissues, but survival log-rank tests did not support prognostic stratification.

**Figure 4. Spatial source-table quantification supports progression and stromal remodeling context.** HRA003627 ROI-level signatures support keratinization/differentiation loss and cancerization/progression gain. HRA008846 Table S3 supports OGT/CCND1 and CAF/ECM source-table rows, while Table S6 lacks direct JAG1-NOTCH1 ligand-receptor evidence.

**Figure 5. Reviewer-enforced claim map for manuscript submission.** The audit map distinguishes supported claims from unsupported survival, direct ligand-receptor and drug-sensitivity claims.

## References

{refs}
"""


def write_references() -> None:
    fields = ["id", "authors", "title", "journal", "year", "volume", "pages", "doi", "pmid"]
    with (REF_DIR / "references.tsv").open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fields, delimiter="\t")
        writer.writeheader()
        for ref in REFERENCES:
            writer.writerow(ref)
    bib_lines = []
    for ref in REFERENCES:
        entry_type = "article"
        bib_lines.append(f"@{entry_type}{{{ref['id']},")
        bib_lines.append(f"  author = {{{ref['authors']}}},")
        bib_lines.append(f"  title = {{{ref['title']}}},")
        bib_lines.append(f"  journal = {{{ref['journal']}}},")
        bib_lines.append(f"  year = {{{ref['year']}}},")
        bib_lines.append(f"  volume = {{{ref['volume']}}},")
        bib_lines.append(f"  pages = {{{ref['pages']}}},")
        if ref.get("doi"):
            bib_lines.append(f"  doi = {{{ref['doi']}}},")
        if ref.get("pmid"):
            bib_lines.append(f"  pmid = {{{ref['pmid']}}},")
        bib_lines.append("}\n")
    (REF_DIR / "references.bib").write_text("\n".join(bib_lines), encoding="utf-8")


def write_markdown_manuscript() -> Path:
    path = MANUSCRIPT_DIR / "ESCC_spatial_source_table_manuscript.md"
    path.write_text(manuscript_text(), encoding="utf-8")
    return path


def add_markdown_to_docx(md_path: Path, docx_path: Path, figure_paths: list[dict[str, str]]) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for line in md_path.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(16)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif re.match(r"^\d+\. ", stripped):
            doc.add_paragraph(stripped, style=None)
        elif stripped.startswith("**Figure "):
            p = doc.add_paragraph()
            run = p.add_run(re.sub(r"\*\*", "", stripped))
            run.bold = True
        else:
            doc.add_paragraph(stripped)
    doc.add_page_break()
    doc.add_heading("Main figures", level=1)
    for entry in figure_paths:
        fig = entry["figure"]
        png = Path(entry["png"])
        doc.add_heading(fig, level=2)
        doc.add_picture(str(png), width=Inches(6.5))
    doc.save(docx_path)


def write_main_tables() -> None:
    dataset_summary_table().to_csv(MANUSCRIPT_DIR / "Table1_dataset_summary.tsv", sep="\t", index=False)
    evidence_gate_table().to_csv(MANUSCRIPT_DIR / "Table2_evidence_gate_summary.tsv", sep="\t", index=False)


def path_status(path: Path) -> tuple[str, int, str]:
    if path.exists():
        return "exists", path.stat().st_size, path.stat().st_mtime_ns and datetime.fromtimestamp(path.stat().st_mtime, UTC).date().isoformat()
    return "not_found", 0, ""


def data_source_manifest_table() -> pd.DataFrame:
    rows: list[dict[str, str | int]] = []

    def add(resource: str, role: str, accession: str, url: str, path: Path, note: str) -> None:
        status, bytes_, modified = path_status(path)
        rows.append(
            {
                "resource": resource,
                "role": role,
                "accession_or_release": accession,
                "url": url,
                "local_path": str(path),
                "bytes": bytes_,
                "local_status": status,
                "download_or_cache_date": modified or RUN_DATE,
                "note": note,
            }
        )

    add(
        "TCGA ESCA squamous subset via UCSC Xena/Toil",
        "Discovery bulk expression",
        "TCGA ESCA; squamous histology subset; Toil recompute resource",
        "https://xenabrowser.net/datapages/",
        TABLE_ROOT / "deep_axis_tcga_immune_associations.tsv",
        "Derived association table; source expression accessed through the upstream workflow.",
    )
    add(
        "GSE47404",
        "Independent tumor-only bulk validation",
        "GSE47404; GPL6480",
        "https://ftp.ncbi.nlm.nih.gov/geo/series/GSE47nnn/GSE47404/matrix/GSE47404_series_matrix.txt.gz",
        OUT_ROOT / "data" / "geo" / "GSE47404" / "GSE47404_series_matrix.txt.gz",
        "Download manifest: deep_axis_geo_download_manifest.tsv.",
    )
    add(
        "GSE53625",
        "Paired tumor-normal and survival audit after target probe rescue",
        "GSE53625; GPL18109 feature-number matrix",
        "https://ftp.ncbi.nlm.nih.gov/geo/series/GSE53nnn/GSE53625/matrix/GSE53625_series_matrix.txt.gz",
        OUT_ROOT / "data" / "geo" / "GSE53625" / "GSE53625_series_matrix.txt.gz",
        "Probe rescue used first raw member plus Ensembl target sequence cache; target-scope only.",
    )
    add(
        "GSE53625 first raw member",
        "Probe-sequence extraction for target rescue",
        "GSM1296956 first RAW member",
        "GEO supplementary raw member",
        OUT_ROOT / "data" / "geo" / "GSE53625" / "GSM1296956_first_raw_member.txt.gz",
        "Used to extract probe sequences; no direct axis gene-symbol hits in original raw annotation.",
    )
    add(
        "Ensembl REST sequence cache",
        "Target sequence lookup for GSE53625 probe rescue",
        "Ensembl REST API; cached on local run date",
        "https://rest.ensembl.org/",
        TABLE_ROOT / "gse53625_ensembl_sequence_index.tsv",
        "Sequence index table records Ensembl gene IDs, transcript counts and executor/reviewer IDs.",
    )
    add(
        "HRA003627 source data",
        "Spatial ROI source-table progression signatures",
        "Nat Commun 2023 source data",
        "https://static-content.springer.com/esm/art%3A10.1038%2Fs41467-023-40343-5/MediaObjects/41467_2023_40343_MOESM4_ESM.xlsx",
        OUT_ROOT / "data" / "open_source_tables" / "HRA003627_NatCommun2023_source_data.xlsx",
        "Open source table manifest records workbook sheet dimensions and download date.",
    )
    add(
        "HRA008846 Table S3",
        "Spatial DEG and perturbation source-table rows",
        "Cell Reports Medicine 2026 supplementary Table S3",
        "https://pmc.ncbi.nlm.nih.gov/articles/instance/13006417/bin/mmc2.xlsx",
        OUT_ROOT / "data" / "open_source_tables" / "HRA008846_TableS3_DEG.xlsx",
        "Valid cached XLSX used; earlier automated attempt logged invalid HTML in open_spatial_source_table_manifest.tsv.",
    )
    add(
        "HRA008846 Table S4",
        "Cell abundance source-table rows",
        "Cell Reports Medicine 2026 supplementary Table S4",
        "https://pmc.ncbi.nlm.nih.gov/articles/instance/13006417/bin/mmc3.xlsx",
        OUT_ROOT / "data" / "open_source_tables" / "HRA008846_TableS4_cell_abundance.xlsx",
        "Valid cached XLSX used; earlier automated attempt logged invalid HTML in open_spatial_source_table_manifest.tsv.",
    )
    add(
        "HRA008846 Table S6",
        "Ligand-receptor source-table rows",
        "Cell Reports Medicine 2026 supplementary Table S6",
        "https://pmc.ncbi.nlm.nih.gov/articles/instance/13006417/bin/mmc4.xlsx",
        OUT_ROOT / "data" / "open_source_tables" / "HRA008846_TableS6_ligand_receptor.xlsx",
        "Valid cached XLSX used; direct JAG1-NOTCH1 rows were absent in parsed axis-relevant rows.",
    )
    add(
        "GDSC2 fitted dose response",
        "Drug target-class context only",
        "GDSC release 8.5; GDSC2 fitted dose response 27Oct23",
        "https://cog.sanger.ac.uk/cancerrxgene/GDSC_release8.5/GDSC2_fitted_dose_response_27Oct23.xlsx",
        OUT_ROOT / "data" / "gdsc" / "GDSC2_fitted_dose_response_27Oct23.xlsx",
        "No axis-expression/drug-response joint model was fitted.",
    )
    return pd.DataFrame(rows)


def supplementary_readme_table() -> pd.DataFrame:
    rows = [
        ("S1_axis_definitions", "Pre-specified axis genes and claim boundaries."),
        ("S2_TCGA_immune", "TCGA ESCC immune/stromal marker associations."),
        ("S3_TCGA_pathway", "TCGA ESCC pathway associations."),
        ("S4_GSE47404_assoc", "Independent GSE47404 expression replication associations."),
        ("S5_GSE47404_clinical", "Limited GSE47404 clinical/pathology associations."),
        ("S6_GSE53625_mapping", "Target-scope GSE53625 probe-sequence mapping."),
        ("S7_GSE53625_validation", "GSE53625 survival and paired tumor-normal validation."),
        ("S8_HRA003627_source", "HRA003627 ROI-level source-table signature quantification."),
        ("S9_HRA008846_DEG", "HRA008846 differential-expression and perturbation source-table hits."),
        ("S10_HRA008846_cells", "HRA008846 cell-abundance source-table trends."),
        ("S11_HRA008846_LR", "HRA008846 ligand-receptor rows and direct JAG1-NOTCH1 flag."),
        ("S12_GDSC2_drugs", "GDSC2 ESCA drug target-class context; not prediction."),
        ("S13_review_gates", "Executor/reviewer gate decisions."),
        ("S14_subagent_audit", "Independent subagent audit summary."),
        ("S15_references", "Reference metadata used in the manuscript."),
        ("S16_data_manifest", "Data source URLs, versions, file sizes and conflict notes."),
    ]
    return pd.DataFrame(rows, columns=["sheet", "description"])


def copy_code_artifacts() -> None:
    script_names = [
        "run_real_workflow.py",
        "run_spatial_axis_deep_validation.py",
        "run_independent_patient_and_spatial_quant.py",
        "build_final_submission_package.py",
    ]
    for name in script_names:
        src = ROOT / "scripts" / name
        if src.exists():
            shutil.copy2(src, CODE_DIR / name)
    package_src = ROOT / "escc_splice_workflow"
    package_dst = CODE_DIR / "escc_splice_workflow"
    if package_dst.exists():
        shutil.rmtree(package_dst)
    if package_src.exists():
        shutil.copytree(package_src, package_dst, ignore=shutil.ignore_patterns("__pycache__", "*.pyc"))
    config_src = ROOT / "project_config.yaml"
    if config_src.exists():
        shutil.copy2(config_src, CODE_DIR / "project_config.yaml")
        shutil.copy2(config_src, PACKAGE_ROOT / "project_config.yaml")
    (CODE_DIR / "requirements.txt").write_text(
        "\n".join(["pandas", "numpy", "openpyxl", "python-docx", "Pillow", "xenaPython"]) + "\n",
        encoding="utf-8",
    )
    (CODE_DIR / "README_code.md").write_text(
        "\n".join(
            [
                "# Reproducibility Code",
                "",
                "These scripts generated the public-data validation tables, manuscript figures and submission package.",
                "",
                "- run_real_workflow.py",
                "- run_spatial_axis_deep_validation.py",
                "- run_independent_patient_and_spatial_quant.py",
                "- build_final_submission_package.py",
                "- escc_splice_workflow/ helper package",
                "- project_config.yaml",
                "- requirements.txt",
                "",
                "Run order: deep validation, independent patient/spatial quantification, then final package build.",
                "The workflow uses public data only and records executor/reviewer gates in output tables.",
            ]
        )
        + "\n",
        encoding="utf-8",
    )


def figure_caption_map() -> dict[str, str]:
    return {
        "Fig 1": (
            "Fig 1. Evidence-gated integrative workflow. Spatially nominated axes were evaluated in bulk ESCC cohorts, "
            "a target-rescued paired patient cohort, open spatial source tables and review gates. The final claim "
            "boundary limits the manuscript to association and source-table reproducibility."
        ),
        "Fig 2": (
            "Fig 2. Cross-cohort axis associations with immune and pathway programs. Spearman correlations show that "
            "the CAF/ECM phenotype is strongly associated with CAF, ECM remodeling and EMT programs in TCGA ESCC and "
            "GSE47404. The OGT/PI3K axis is most consistently associated with hypoxia. Asterisks indicate FDR<0.05."
        ),
        "Fig 3": (
            "Fig 3. GSE53625 target probe rescue supports tumor-normal upshift but not survival prediction. Target-scope "
            "probe rescue recovered axis gene coverage. Both axes were elevated in paired tumors relative to adjacent "
            "normal tissues, but survival log-rank tests did not support prognostic stratification."
        ),
        "Fig 4": (
            "Fig 4. Spatial source-table quantification supports progression and stromal remodeling context. HRA003627 "
            "ROI-level signatures support keratinization/differentiation loss and cancerization/progression gain. "
            "HRA008846 Table S3 supports OGT/CCND1 and CAF/ECM source-table rows, while Table S6 lacks direct "
            "JAG1-NOTCH1 ligand-receptor evidence."
        ),
        "Fig 5": (
            "Fig 5. Reviewer-enforced claim map for manuscript submission. The audit map distinguishes supported claims "
            "from unsupported survival, direct ligand-receptor and drug-sensitivity claims."
        ),
    }


def table_to_markdown(df: pd.DataFrame) -> str:
    columns = [str(col) for col in df.columns]
    lines = [
        "| " + " | ".join(columns) + " |",
        "| " + " | ".join(["---"] * len(columns)) + " |",
    ]
    for _, row in df.iterrows():
        vals = [str(row[col]).replace("\n", " ") for col in df.columns]
        lines.append("| " + " | ".join(vals) + " |")
    return "\n".join(lines)


def plos_one_text() -> str:
    text = manuscript_text()
    plos_abstract = (
        "Esophageal squamous cell carcinoma (ESCC) progression involves tumor-intrinsic and stromal programs, but "
        "spatial observations often require cautious validation in public cohorts. We evaluated two pre-specified "
        "spatially informed axes using public transcriptomic cohorts and open spatial source tables. TCGA ESCC and "
        "GSE47404 supported a reproducible CAF/ECM stromal-remodeling phenotype, with strong correlations to CAF, "
        "ECM remodeling and EMT programs. The OGT/PI3K/TLS axis showed its most stable association with hypoxia "
        "rather than with a specific TLS signal. Target-scope probe rescue in GSE53625 recovered axis gene coverage "
        "and showed that both axes were elevated in 179 paired tumors relative to adjacent normal tissue; however, "
        "survival audits did not support prognostic stratification. HRA003627 source tables supported disease-stage "
        "progression signatures, and HRA008846 source tables supported OGT/CCND1 and CAF/ECM context but contained "
        "no direct JAG1-NOTCH1 ligand-receptor rows. GDSC2 summaries were retained only as target-class context. "
        "Overall, this study supports an association-focused public-data manuscript centered on a reproducible "
        "CAF/ECM stromal-remodeling phenotype in ESCC, while explicitly rejecting independent survival prediction, "
        "direct JAG1-NOTCH1 signaling and drug-sensitivity prediction claims."
    )
    text = re.sub(r"## Abstract\n\n.*?\n\nKeywords:", f"## Abstract\n\n{plos_abstract}\n\nKeywords:", text, flags=re.S)
    prefix, rest = text.split("\n## Figure legends\n", 1)
    _, refs = rest.split("\n## References\n", 1)
    prefix = prefix.replace("(Figure 2)", "(Fig 2)")
    prefix = prefix.replace("(Figure 3)", "(Fig 3)")
    prefix = prefix.replace("(Figure 4)", "(Fig 4)")
    prefix = prefix.replace(
        "A pre-submission evidence gate restricted the manuscript to association and source-table reproducibility claims, not prognostic biomarker, causal mechanism or therapeutic-response claims.",
        "A pre-submission evidence gate restricted the manuscript to association and source-table reproducibility claims, not prognostic biomarker, causal mechanism or therapeutic-response claims. Dataset scope is summarized in Table 1.\n\n[[TABLE1]]\n\nThe complete evidence-gated workflow is summarized in Fig 1.\n\n"
        + figure_caption_map()["Fig 1"],
    )
    prefix = prefix.replace(
        "These cross-cohort associations nominate a reusable CAF/ECM stromal-remodeling phenotype in ESCC (Fig 2).",
        "These cross-cohort associations nominate a reusable CAF/ECM stromal-remodeling phenotype in ESCC (Fig 2).\n\n"
        + figure_caption_map()["Fig 2"],
    )
    prefix = prefix.replace(
        "Thus, GSE53625 supports tumor-normal expression shifts but not independent survival prediction (Fig 3).",
        "Thus, GSE53625 supports tumor-normal expression shifts but not independent survival prediction (Fig 3).\n\n"
        + figure_caption_map()["Fig 3"],
    )
    prefix = prefix.replace(
        "Therefore, HRA008846 supports stromal remodeling and epithelial progression context but not a proven JAG1-NOTCH1 cell-cell communication mechanism (Fig 4).",
        "Therefore, HRA008846 supports stromal remodeling and epithelial progression context but not a proven JAG1-NOTCH1 cell-cell communication mechanism (Fig 4).\n\n"
        + figure_caption_map()["Fig 4"],
    )
    prefix = prefix.replace(
        "Finally, the current evidence does not support survival prediction, direct JAG1-NOTCH1 signaling or therapeutic-response prediction.",
        "Finally, the current evidence does not support survival prediction, direct JAG1-NOTCH1 signaling or therapeutic-response prediction. Evidence-gate decisions are summarized in Table 2.\n\n[[TABLE2]]\n\nThe claim-control map is shown in Fig 5.\n\n"
        + figure_caption_map()["Fig 5"],
    )
    supporting = (
        "\n## Supporting information captions\n\n"
        "S1 Table. Supplementary tables for axis definitions, cohort associations, GSE53625 probe rescue, "
        "spatial source-table quantification, GDSC2 target-class context, review gates, references and data-source manifest.\n\n"
        "S2 References. Machine-readable reference metadata used to build the manuscript bibliography.\n\n"
        "S3 Internal review gates. Executor/reviewer claim-control decisions used to restrict unsupported survival, "
        "mechanistic and drug-response claims.\n"
    )
    return prefix + "\n## References\n" + refs + supporting


def write_plos_one_markdown() -> Path:
    text = plos_one_text()
    text = text.replace("[[TABLE1]]", "**Table 1. Dataset and resource summary.**\n\n" + table_to_markdown(dataset_summary_table()))
    text = text.replace("[[TABLE2]]", "**Table 2. Evidence-gate summary.**\n\n" + table_to_markdown(evidence_gate_table()))
    path = PLOS_MANUSCRIPT_DIR / "ESCC_spatial_source_table_PLOS_ONE_main_text.md"
    path.write_text(text, encoding="utf-8")
    return path


def add_dataframe_table(doc: Document, title: str, df: pd.DataFrame, note: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, col in enumerate(df.columns):
        hdr[idx].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(df.columns):
            cells[idx].text = str(row[col])
    p = doc.add_paragraph(note)
    p.runs[0].italic = True


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_text, fld_end])


def apply_plos_docx_format(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(0)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        sect_pr = section._sectPr
        for existing in sect_pr.findall(qn("w:lnNumType")):
            sect_pr.remove(existing)
        ln_num = OxmlElement("w:lnNumType")
        ln_num.set(qn("w:countBy"), "1")
        ln_num.set(qn("w:restart"), "continuous")
        sect_pr.append(ln_num)
        add_page_number(section.footer.paragraphs[0])


def add_plos_text_to_docx(docx_path: Path, inline_figures: dict[str, Path] | None = None) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    lines = plos_one_text().splitlines()
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped == "[[TABLE1]]":
            add_dataframe_table(
                doc,
                "Table 1. Dataset and resource summary.",
                dataset_summary_table(),
                "Note: all data were public and de-identified; claim boundaries are enforced by review gates.",
            )
            continue
        if stripped == "[[TABLE2]]":
            add_dataframe_table(
                doc,
                "Table 2. Evidence-gate summary.",
                evidence_gate_table(),
                "Note: PASS WITH LIMITS indicates association or source-table support without prognostic or mechanistic proof.",
            )
            continue
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(16)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif re.match(r"^\d+\. ", stripped):
            doc.add_paragraph(stripped)
        elif stripped.startswith("Fig "):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            match = re.match(r"^(Fig\s+\d+)\.", stripped)
            if inline_figures and match:
                figure_path = inline_figures.get(match.group(1))
                if figure_path and figure_path.exists():
                    fig_para = doc.add_paragraph()
                    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    fig_para.add_run().add_picture(str(figure_path), width=Inches(6.5))
        else:
            doc.add_paragraph(stripped)
    apply_plos_docx_format(doc)
    doc.save(docx_path)


def nonwhite_bbox(image: Image.Image) -> tuple[int, int, int, int] | None:
    rgb = image.convert("RGB")
    px = rgb.load()
    width, height = rgb.size
    xs: list[int] = []
    ys: list[int] = []
    for y in range(0, height, max(1, height // 600)):
        for x in range(0, width, max(1, width // 600)):
            r, g, b = px[x, y]
            if not (r > 248 and g > 248 and b > 248):
                xs.append(x)
                ys.append(y)
    if not xs or not ys:
        return None
    return min(xs), min(ys), max(xs), max(ys)


def create_plos_one_figures(figure_entries: list[dict[str, str]]) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for idx, entry in enumerate(figure_entries, start=1):
        source = Path(entry["png"])
        image = Image.open(source).convert("RGB")
        width, height = image.size
        crop_top = 0
        cropped = image.crop((0, crop_top, width, height))
        if cropped.size[0] > 2250:
            new_height = int(cropped.size[1] * 2250 / cropped.size[0])
            cropped = cropped.resize((2250, new_height), Image.Resampling.LANCZOS)
        clean_png = PLOS_FIGURE_DIR / f"Fig{idx}.png"
        tif = PLOS_FIGURE_DIR / f"Fig{idx}.tif"
        cropped.save(clean_png, dpi=(300, 300))
        cropped.save(tif, dpi=(300, 300), compression="tiff_lzw")
        bbox = nonwhite_bbox(cropped)
        margin = None
        if bbox:
            x1, y1, x2, y2 = bbox
            margin = min(x1, y1, cropped.size[0] - x2 - 1, cropped.size[1] - y2 - 1)
        records.append(
            {
                "figure": f"Fig {idx}",
                "source_png": str(source.relative_to(PACKAGE_ROOT)),
                "plos_png": str(clean_png.relative_to(PACKAGE_ROOT)),
                "plos_tif": str(tif.relative_to(PACKAGE_ROOT)),
                "width_px": cropped.size[0],
                "height_px": cropped.size[1],
                "dpi": "300",
                "title_removed_in_source": "yes",
                "title_removed_by_crop_px": crop_top,
                "nonwhite_bbox": "" if bbox is None else ",".join(map(str, bbox)),
                "min_nonwhite_margin_px": "" if margin is None else margin,
                "technical_status": "pass" if cropped.size[0] >= 2250 and cropped.size[1] > 900 and margin is not None and margin >= 0 else "review",
            }
        )
    audit = pd.DataFrame(records)
    audit.to_csv(PLOS_REVIEW_DIR / "figure_technical_audit.tsv", sep="\t", index=False)
    return records


def write_plos_one_compliance_matrix(plos_figures: list[dict[str, Any]]) -> Path:
    rows = [
        {
            "requirement": "Main manuscript should not include figures; figures submitted as individual files",
            "source": "PLOS ONE submission guidelines",
            "status": "pass",
            "evidence": "PLOS ONE DOCX contains text, captions and editable tables only; Fig1-Fig5 are separate TIFF files.",
            "action": "Use plos_one_submission/manuscript and plos_one_submission/figures for submission.",
        },
        {
            "requirement": "Figure captions in manuscript text after first citation, using Fig labels",
            "source": "PLOS ONE figure guidelines",
            "status": "pass",
            "evidence": "PLOS ONE markdown/DOCX inserts Fig 1-Fig 5 captions after first in-text citations.",
            "action": "Do not upload captions as separate files.",
        },
        {
            "requirement": "Tables included in main manuscript as editable tables after first citation",
            "source": "PLOS ONE submission guidelines",
            "status": "pass",
            "evidence": "Table 1 and Table 2 are embedded as Word tables in the PLOS ONE DOCX.",
            "action": "Keep Table1/Table2 in the main text; large source tables remain supporting information.",
        },
        {
            "requirement": "References numbered in order and complete",
            "source": "PLOS ONE submission guidelines",
            "status": "pass_with_author_check",
            "evidence": "References are numbered and include DOI/PMID where available; authors should confirm no cited work has been retracted before submission.",
            "action": "Run final reference-manager or PubMed retraction check immediately before upload.",
        },
        {
            "requirement": "Figures at publication-quality resolution",
            "source": "PLOS ONE figure guidelines",
            "status": "pass",
            "evidence": f"{len(plos_figures)} TIFF figures generated at 300 dpi; width normalized to full-page target where needed.",
            "action": "Optional: run PLOS NAAS before final upload.",
        },
        {
            "requirement": "Statistical methods sufficiently detailed and software/code identified",
            "source": "PLOS ONE statistical reporting guidance",
            "status": "revise_minor",
            "evidence": "Methods report tests and FDR; code path is listed, but exact Python/R package versions are not yet in the manuscript.",
            "action": "Add package/version table if journal asks during technical check.",
        },
        {
            "requirement": "Double spacing, page numbers and continuous line numbers",
            "source": "PLOS ONE submission guidelines",
            "status": "pass_with_render_limit",
            "evidence": "PLOS ONE DOCX generator applies double spacing, 1-inch margins, footer PAGE field and continuous line numbering XML.",
            "action": "Open in Word before upload to confirm field rendering.",
        },
        {
            "requirement": "Complete title page, author details, funding and competing interests",
            "source": "PLOS ONE submission guidelines",
            "status": "requires_author_input",
            "evidence": "Scientific content package cannot infer real author names, affiliations, funding or conflicts.",
            "action": "Authors must replace placeholders before actual journal upload.",
        },
        {
            "requirement": "No unsupported clinical or mechanistic overclaiming",
            "source": "Internal scientific review gate",
            "status": "pass",
            "evidence": "Manuscript explicitly rejects survival prediction, direct JAG1-NOTCH1 mechanism and drug-sensitivity prediction.",
            "action": "Maintain conservative title and abstract language.",
        },
    ]
    path = PLOS_REVIEW_DIR / "plos_one_compliance_matrix.tsv"
    pd.DataFrame(rows).to_csv(path, sep="\t", index=False)
    return path


def write_plos_one_package(figure_entries: list[dict[str, str]], supp_path: Path) -> dict[str, Any]:
    md_path = write_plos_one_markdown()
    plos_figures = create_plos_one_figures(figure_entries)
    docx_path = PLOS_MANUSCRIPT_DIR / "ESCC_spatial_source_table_PLOS_ONE_main_text.docx"
    add_plos_text_to_docx(docx_path)
    inline_docx_path = PLOS_MANUSCRIPT_DIR / "ESCC_spatial_source_table_PLOS_ONE_main_text_with_inline_figures.docx"
    inline_figure_map = {
        item["figure"]: PACKAGE_ROOT / item["plos_png"]
        for item in plos_figures
    }
    add_plos_text_to_docx(inline_docx_path, inline_figures=inline_figure_map)
    shutil.copy2(supp_path, PLOS_SUPPORTING_DIR / "S1_Table.xlsx")
    shutil.copy2(REF_DIR / "references.tsv", PLOS_SUPPORTING_DIR / "S2_References.tsv")
    review_source = REVIEW_ROOT / "independent_patient_and_spatial_quant_review.tsv"
    if review_source.exists():
        shutil.copy2(review_source, PLOS_SUPPORTING_DIR / "S3_Internal_review_gates.tsv")
    compliance_path = write_plos_one_compliance_matrix(plos_figures)
    readme = PLOS_DIR / "README_PLOS_ONE_submission.md"
    readme.write_text(
        "\n".join(
            [
                "# PLOS ONE Target-Journal Submission Subpackage",
                "",
                "Use this subpackage for PLOS ONE first submission.",
                "",
                "## Files",
                "",
                "- manuscript/ESCC_spatial_source_table_PLOS_ONE_main_text.docx",
                "- manuscript/ESCC_spatial_source_table_PLOS_ONE_main_text_with_inline_figures.docx (review copy; not for PLOS figure upload)",
                "- manuscript/ESCC_spatial_source_table_PLOS_ONE_main_text.md",
                "- figures/Fig1.tif through figures/Fig5.tif",
                "- supporting_information/S1_Table.xlsx",
                "- internal_review/plos_one_compliance_matrix.tsv",
                "- internal_review/figure_technical_audit.tsv",
                "",
                "## Claim limits retained",
                "",
                "- No independent survival prediction claim.",
                "- No direct JAG1-NOTCH1 ligand-receptor mechanism claim.",
                "- No drug-sensitivity prediction claim.",
                "- No raw spatial matrix reanalysis claim.",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    return {
        "plos_markdown": str(md_path),
        "plos_docx": str(docx_path),
        "plos_inline_review_docx": str(inline_docx_path),
        "plos_figures": plos_figures,
        "plos_compliance": str(compliance_path),
    }


def sanitize_sheet_name(name: str) -> str:
    cleaned = re.sub(r"[\[\]\:\*\?\/\\]", "_", name)
    return cleaned[:31]


def copy_tsv(src: Path, dst: Path) -> None:
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)


def write_supplementary_tables() -> Path:
    table_specs = [
        ("README", supplementary_readme_table()),
        ("S1_axis_definitions", pd.DataFrame([
            {"axis_id": "ogt_pi3k_tls_axis", "axis_label": "OGT/PI3K/TLS axis", "genes": "OGT,PIK3CA,AKT1,CCND1,LAMB1,SPP1,KRT17,APOBEC3A,JAG1,NOTCH1", "claim_boundary": "hypoxia/epithelial progression context"},
            {"axis_id": "caf_epi_jag1_notch_niche", "axis_label": "CAF/ECM stromal-remodeling phenotype", "genes": "JAG1,NOTCH1,FAP,COL1A1,COL1A2,POSTN,CXCL1,CXCL8,SPP1", "claim_boundary": "CAF/ECM phenotype; not direct JAG1-NOTCH1 proof"},
        ])),
        ("S2_TCGA_immune", read_tsv(TABLE_ROOT / "deep_axis_tcga_immune_associations.tsv")),
        ("S3_TCGA_pathway", read_tsv(TABLE_ROOT / "deep_axis_tcga_pathway_associations.tsv")),
        ("S4_GSE47404_assoc", read_tsv(TABLE_ROOT / "deep_axis_geo_gse47404_associations.tsv")),
        ("S5_GSE47404_clinical", read_tsv(TABLE_ROOT / "deep_axis_geo_gse47404_clinical_associations.tsv")),
        ("S6_GSE53625_mapping", read_tsv(TABLE_ROOT / "gse53625_probe_sequence_mapping.tsv")),
        ("S7_GSE53625_validation", pd.concat([
            read_tsv(TABLE_ROOT / "gse53625_rescue_survival_validation.tsv"),
            read_tsv(TABLE_ROOT / "gse53625_rescue_tumor_normal_validation.tsv"),
        ], ignore_index=True, sort=False)),
        ("S8_HRA003627_source", read_tsv(TABLE_ROOT / "hra003627_source_table_quantification.tsv")),
        ("S9_HRA008846_DEG", read_tsv(TABLE_ROOT / "hra008846_deg_axis_hits.tsv")),
        ("S10_HRA008846_cells", read_tsv(TABLE_ROOT / "hra008846_cell_abundance_trends.tsv")),
        ("S11_HRA008846_LR", read_tsv(TABLE_ROOT / "hra008846_ligand_receptor_axis_hits.tsv")),
        ("S12_GDSC2_drugs", read_tsv(TABLE_ROOT / "deep_axis_gdsc2_esca_drug_response.tsv")),
        ("S13_review_gates", read_tsv(REVIEW_ROOT / "independent_patient_and_spatial_quant_review.tsv")),
        ("S14_subagent_audit", read_tsv(REVIEW_ROOT / "independent_patient_and_spatial_quant_subagent_audit.tsv")),
        ("S15_references", pd.DataFrame(REFERENCES)),
        ("S16_data_manifest", data_source_manifest_table()),
    ]
    workbook_path = SUPP_DIR / "Supplementary_Tables_ESCC_spatial_source_table.xlsx"
    with pd.ExcelWriter(workbook_path, engine="openpyxl") as writer:
        for sheet_name, df in table_specs:
            df.to_excel(writer, sheet_name=sanitize_sheet_name(sheet_name), index=False)
    wb = load_workbook(workbook_path)
    for ws in wb.worksheets:
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="4F81BD")
        for col_cells in ws.columns:
            max_len = min(max(len(str(cell.value or "")) for cell in col_cells), 70)
            ws.column_dimensions[col_cells[0].column_letter].width = max(12, max_len + 2)
    wb.save(workbook_path)
    # Machine-readable TSV copies.
    for sheet_name, df in table_specs:
        df.to_csv(SUPP_DIR / f"{sheet_name}.tsv", sep="\t", index=False)
    return workbook_path


def write_cover_letter() -> Path:
    text = f"""Dear Editor,

We are pleased to submit the manuscript entitled "{TITLE}" for consideration as an original research article.

This manuscript is intentionally positioned as a public transcriptomic and spatial source-table reproducibility study in ESCC. It integrates TCGA ESCC, GSE47404, a target-rescued GSE53625 paired cohort, HRA003627 source data, HRA008846 supplementary source tables and GDSC2 target-class context. The main finding is a reproducible CAF/ECM stromal-remodeling phenotype, with an auxiliary OGT/PI3K-associated hypoxia and epithelial progression context.

We have explicitly avoided unsupported claims. The manuscript does not claim independent survival prediction, direct CAF-epithelial JAG1-NOTCH1 ligand-receptor signaling, raw spatial matrix reanalysis or drug-sensitivity prediction. These boundaries are documented in the reviewer-audit tables included with the submission package.

All data are public and all scripts and source tables needed to reproduce the figures and supplementary tables are included in the package.

Sincerely,

The authors
"""
    path = MANUSCRIPT_DIR / "cover_letter.md"
    path.write_text(text, encoding="utf-8")
    return path


def write_readme(figure_entries: list[dict[str, str]], supp_path: Path) -> None:
    lines = [
        "# ESCC Spatial Source-Table Submission Package",
        "",
        f"Build date: {RUN_DATE}",
        "",
        "## Contents",
        "",
        "- manuscript/ESCC_spatial_source_table_manuscript.md",
        "- manuscript/ESCC_spatial_source_table_manuscript.docx",
        "- manuscript/Table1_dataset_summary.tsv",
        "- manuscript/Table2_evidence_gate_summary.tsv",
        "- figures/Figure1-Figure5 PNG and SVG files",
        f"- supplementary_tables/{supp_path.name}",
        "- supplementary_tables/S*.tsv machine-readable copies",
        "- references/references.tsv and references.bib",
        "- review_audit/*.tsv",
        "- plos_one_submission/ target-journal manuscript, figures, supporting information and compliance checks",
        "",
        "## Claim limits",
        "",
        "- Do not claim independent survival prediction.",
        "- Do not claim direct CAF-epithelial JAG1-NOTCH1 ligand-receptor signaling.",
        "- Do not claim drug-sensitivity prediction.",
        "- Do not describe source-table quantification as raw spatial matrix reanalysis.",
        "",
        "## Figure files",
        "",
    ]
    for entry in figure_entries:
        lines.append(f"- {entry['figure']}: {Path(entry['png']).name}; {Path(entry['svg']).name}")
    (PACKAGE_ROOT / "README_submission_package.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def copy_review_artifacts() -> None:
    for path in [
        REVIEW_ROOT / "independent_patient_and_spatial_quant_review.tsv",
        REVIEW_ROOT / "independent_patient_and_spatial_quant_subagent_audit.tsv",
        REVIEW_ROOT / "spatial_axis_manuscript_signoff.tsv",
    ]:
        if path.exists():
            shutil.copy2(path, REVIEW_DIR / path.name)


def write_repro_check(figure_entries: list[dict[str, str]], manuscript_md: Path, manuscript_docx: Path, supp_path: Path) -> Path:
    checks = []
    targets = [manuscript_md, manuscript_docx, supp_path, REF_DIR / "references.tsv", REF_DIR / "references.bib"]
    targets += [Path(entry["png"]) for entry in figure_entries]
    targets += [Path(entry["svg"]) for entry in figure_entries]
    targets += [
        PLOS_MANUSCRIPT_DIR / "ESCC_spatial_source_table_PLOS_ONE_main_text.md",
        PLOS_MANUSCRIPT_DIR / "ESCC_spatial_source_table_PLOS_ONE_main_text.docx",
        PLOS_MANUSCRIPT_DIR / "ESCC_spatial_source_table_PLOS_ONE_main_text_with_inline_figures.docx",
        PLOS_SUPPORTING_DIR / "S1_Table.xlsx",
        PLOS_REVIEW_DIR / "plos_one_compliance_matrix.tsv",
        PLOS_REVIEW_DIR / "figure_technical_audit.tsv",
        CODE_DIR / "run_spatial_axis_deep_validation.py",
        CODE_DIR / "run_independent_patient_and_spatial_quant.py",
        CODE_DIR / "run_real_workflow.py",
        CODE_DIR / "build_final_submission_package.py",
        CODE_DIR / "README_code.md",
        CODE_DIR / "requirements.txt",
        CODE_DIR / "project_config.yaml",
        PACKAGE_ROOT / "project_config.yaml",
        CODE_DIR / "escc_splice_workflow" / "__init__.py",
    ]
    targets += sorted(PLOS_FIGURE_DIR.glob("Fig*.tif"))
    for target in targets:
        checks.append(
            {
                "artifact": str(target.relative_to(PACKAGE_ROOT)),
                "exists": target.exists(),
                "bytes": target.stat().st_size if target.exists() else 0,
                "status": "pass" if target.exists() and target.stat().st_size > 0 else "fail",
            }
        )
    path = PACKAGE_ROOT / "reproducibility_check.tsv"
    pd.DataFrame(checks).to_csv(path, sep="\t", index=False)
    return path


def zip_package() -> Path:
    zip_path = DELIVERABLE_ROOT / f"ESCC_spatial_source_table_submission_package_{RUN_DATE}.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in PACKAGE_ROOT.rglob("*"):
            if path.is_file():
                zf.write(path, path.relative_to(PACKAGE_ROOT.parent))
    return zip_path


def main() -> None:
    ensure_dirs()
    figure_entries = make_all_figures()
    write_main_tables()
    write_references()
    manuscript_md = write_markdown_manuscript()
    manuscript_docx = MANUSCRIPT_DIR / "ESCC_spatial_source_table_manuscript.docx"
    add_markdown_to_docx(manuscript_md, manuscript_docx, figure_entries)
    supp_path = write_supplementary_tables()
    plos_package = write_plos_one_package(figure_entries, supp_path)
    write_cover_letter()
    copy_review_artifacts()
    copy_code_artifacts()
    write_readme(figure_entries, supp_path)
    repro_path = write_repro_check(figure_entries, manuscript_md, manuscript_docx, supp_path)
    zip_path = zip_package()
    print(
        json.dumps(
            {
                "status": "completed",
                "package_root": str(PACKAGE_ROOT),
                "zip_path": str(zip_path),
                "manuscript_md": str(manuscript_md),
                "manuscript_docx": str(manuscript_docx),
                "supplementary_workbook": str(supp_path),
                "reproducibility_check": str(repro_path),
                "figures": figure_entries,
                "plos_one_submission": plos_package,
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
